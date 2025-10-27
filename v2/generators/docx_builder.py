"""
DOCX Document Builder - Total Health Conferencing
Creates professionally formatted Word documents
"""
from io import BytesIO
from typing import List, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from generators.base import DocumentBuilder
from config.settings import DOCUMENT_SETTINGS


class DOCXBuilder(DocumentBuilder):
    """
    Builds DOCX documents with Total Health branding

    Features:
    - Logo and signature embedding
    - Professional formatting (Times New Roman, proper spacing)
    - Bullet lists with proper indentation
    - Section headers
    """

    def __init__(self):
        """Initialize DOCX builder"""
        self.doc = Document()
        self.default_font = DOCUMENT_SETTINGS["default_font"]
        self.default_font_size = DOCUMENT_SETTINGS["default_font_size"]

    def add_logo(self, logo_bytes: Optional[bytes]):
        """
        Add logo to document (centered, 2.5 inches width)

        Args:
            logo_bytes: Logo image as bytes
        """
        if logo_bytes:
            try:
                paragraph = self.doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(
                    BytesIO(logo_bytes),
                    width=Inches(DOCUMENT_SETTINGS["logo_width_inches"])
                )
            except Exception as e:
                # If image fails, continue without logo
                print(f"Warning: Could not add logo: {e}")

    def add_title(self, title: str, bold: bool = True, centered: bool = True):
        """
        Add title to document

        Args:
            title: Title text
            bold: Make title bold
            centered: Center align title
        """
        paragraph = self.doc.add_paragraph()

        if centered:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run(title)
        run.font.name = self.default_font
        run.font.size = Pt(self.default_font_size)

        if bold:
            run.bold = True

    def add_paragraph(self, text: str, spacing_after: int = 10):
        """
        Add paragraph to document

        Args:
            text: Paragraph text
            spacing_after: Spacing after paragraph (in points)
        """
        paragraph = self.doc.add_paragraph(text)
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.name = self.default_font
        run.font.size = Pt(self.default_font_size)

        # Set spacing
        paragraph.paragraph_format.space_after = Pt(spacing_after)

    def add_section_header(self, header: str, bold: bool = True):
        """
        Add section header to document

        Args:
            header: Header text
            bold: Make header bold
        """
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(header)
        run.font.name = self.default_font
        run.font.size = Pt(self.default_font_size)
        run.bold = bold

        # Spacing
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(8)

    def add_bullet_list(self, items: List[str], indent: float = 0.25):
        """
        Add bullet list to document

        Args:
            items: List of bullet point texts
            indent: Indentation in inches
        """
        for item in items:
            paragraph = self.doc.add_paragraph(item, style='List Bullet')
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.font.name = self.default_font
            run.font.size = Pt(self.default_font_size)

            # Set indentation
            paragraph.paragraph_format.left_indent = Inches(indent)
            paragraph.paragraph_format.space_after = Pt(6)

    def add_signature(self, signature_bytes: Optional[bytes]):
        """
        Add signature image to document (2.0 inches width)

        Args:
            signature_bytes: Signature image as bytes
        """
        if signature_bytes:
            try:
                paragraph = self.doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(
                    BytesIO(signature_bytes),
                    width=Inches(DOCUMENT_SETTINGS["signature_width_inches"])
                )
                paragraph.paragraph_format.space_before = Pt(12)
                paragraph.paragraph_format.space_after = Pt(6)
            except Exception as e:
                print(f"Warning: Could not add signature: {e}")

    def add_contact_info(self, lines: List[str]):
        """
        Add contact information (no extra spacing)

        Args:
            lines: List of contact info lines
        """
        for line in lines:
            paragraph = self.doc.add_paragraph(line)
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.font.name = self.default_font
            run.font.size = Pt(self.default_font_size)
            paragraph.paragraph_format.space_after = Pt(0)

    def add_loa_field(self, label: str, value: str = "", underline_length: int = 30):
        """
        Add LOA signature field (e.g., "Name: _______________")

        Args:
            label: Field label
            value: Field value (if provided)
            underline_length: Length of underline for blank fields
        """
        paragraph = self.doc.add_paragraph()

        # Add label
        run = paragraph.add_run(f"{label}: ")
        run.font.name = self.default_font
        run.font.size = Pt(self.default_font_size)

        # Add value or underline
        if value:
            run = paragraph.add_run(value)
            run.font.name = self.default_font
            run.font.size = Pt(self.default_font_size)
        else:
            run = paragraph.add_run("_" * underline_length)
            run.font.name = self.default_font
            run.font.size = Pt(self.default_font_size)

        paragraph.paragraph_format.space_after = Pt(4)

    def add_numbered_section(self, number: str, title: str):
        """
        Add numbered section header (e.g., "1. Purpose")

        Args:
            number: Section number
            title: Section title
        """
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(f"{number}. {title}")
        run.font.name = self.default_font
        run.font.size = Pt(self.default_font_size)
        run.bold = True

        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(8)

    def add_blank_line(self):
        """Add blank line (paragraph with no text)"""
        self.doc.add_paragraph()

    def get_buffer(self) -> BytesIO:
        """
        Get document as BytesIO buffer

        Returns:
            BytesIO buffer containing DOCX file
        """
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer


# ============================================================================
# CONVENIENCE FUNCTIONS
# ============================================================================

def create_docx_from_paragraphs(
    paragraphs: List[str],
    logo_bytes: Optional[bytes] = None,
    signature_bytes: Optional[bytes] = None,
    contact_lines: Optional[List[str]] = None,
    title: Optional[str] = None
) -> BytesIO:
    """
    Create simple DOCX from paragraphs (for basic documents)

    Args:
        paragraphs: List of paragraph texts
        logo_bytes: Optional logo image
        signature_bytes: Optional signature image
        contact_lines: Optional contact information
        title: Optional document title

    Returns:
        BytesIO buffer containing DOCX file
    """
    builder = DOCXBuilder()

    # Add logo
    if logo_bytes:
        builder.add_logo(logo_bytes)

    # Add title
    if title:
        builder.add_title(title)
        builder.add_blank_line()

    # Add paragraphs
    for para in paragraphs:
        builder.add_paragraph(para)

    # Add signature
    if signature_bytes:
        builder.add_signature(signature_bytes)

    # Add contact info
    if contact_lines:
        builder.add_contact_info(contact_lines)

    return builder.get_buffer()

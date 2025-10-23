# lor_generator_streamlit_app_v_1_pricing_override.py
# LOR/LOA Generator — Total Health Oncology
# Professional Healthcare Document Generation System
# Version 46.0 - Complete Rebuild

import datetime as dt
import os
import base64
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import List, Dict, Optional

import streamlit as st

# --- Password Protection ---
def check_password():
    """Returns `True` if the user entered the correct password."""
    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if st.session_state["password"] == st.secrets["password"]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # Don't store password.
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        # First run, show input for password.
        st.text_input(
            "Password", type="password", on_change=password_entered, key="password"
        )
        return False
    elif not st.session_state["password_correct"]:
        # Password not correct, show input + error.
        st.text_input(
            "Password", type="password", on_change=password_entered, key="password"
        )
        st.error("😕 Password incorrect")
        return False
    else:
        # Password correct.
        return True

if not check_password():
    st.stop()
# --- End Password Protection ---

# Optional dependencies
DOCX_AVAILABLE = False
PDF_AVAILABLE = False
PIL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except Exception:
    pass

try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
        Table, TableStyle
    )
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except Exception:
    pass

try:
    from PIL import Image
    PIL_AVAILABLE = True
except Exception:
    pass

# ---------------------- Configuration & Data ----------------------

# Booth pricing
BOOTH_PRICES = {
    "standard_1d": 5000,
    "standard_2d": 7500,
    "platinum": 10000,
    "best_of": 10000,
    "premier": 15000,
}

# Add-ons pricing (2025 vs 2026)
ADD_ONS_2025 = {
    "program_ad_full": {"label": "Program Guide Full Page Ad", "price": 2000},
    "charging_stations": {"label": "In-Person Charging Station", "price": 2000},
    "wifi_sponsorship": {"label": "Wi-Fi Network Sponsorship", "price": 3000},
    "platform_banner": {"label": "Platform Banner Ad", "price": 2000},
    "email_banner": {"label": "Email Banner Ad", "price": 2500},
    "registration_banner": {"label": "Registration Banner Ad", "price": 2000},
    "networking_reception": {"label": "In-Person Networking Reception", "price": 3500},
    "networking_activity": {"label": "Networking Activity / Excursion", "price": 3500},
    "advisory_board": {"label": "Advisory Board (3-hour)", "price": 30000},
    "non_cme_session": {"label": "Non-CME/CE Session (45 min)", "price": 50000},
}

ADD_ONS_2026 = {
    "program_ad_full": {"label": "Program Guide Full Page Ad", "price": 2000},
    "charging_stations": {"label": "In-Person Charging Station", "price": 3000},
    "wifi_sponsorship": {"label": "Wi-Fi Network Sponsorship", "price": 3000},
    "platform_banner": {"label": "Platform Banner Ad", "price": 2000},
    "email_banner": {"label": "Email Banner Ad", "price": 2500},
    "registration_banner": {"label": "Registration Banner Ad", "price": 2000},
    "networking_reception": {"label": "In-Person Networking Reception", "price": 3500},
    "networking_activity": {"label": "Networking Activity / Excursion", "price": 3500},
    "advisory_board": {"label": "Advisory Board (3-hour)", "price": 30000},
    "non_cme_session": {"label": "Non-CME/CE Session (45 min)", "price": 50000},
}

# Function to get pricing based on event year
def get_add_ons_pricing(event_year):
    """Return the appropriate pricing structure based on event year"""
    if event_year == 2026:
        return ADD_ONS_2026
    else:
        return ADD_ONS_2025

# Add-on descriptions
ADD_ON_BULLETS: Dict[str, List[str]] = {
    "program_ad_full": ["Full-page advertisement in the printed/digital program guide."],
    "charging_stations": [
        "One branded charging station with company artwork.",
        "Includes (1) in-person company representative badge."
    ],
    "wifi_sponsorship": [
        "Exclusive Wi-Fi sponsorship with company logo/name on the Wi-Fi page.",
        "Includes (1) in-person company representative badge."
    ],
    "platform_banner": ["Banner advertisement on the event's digital platform/lobby page."],
    "email_banner": ["Banner placement in a national call-to-action email."],
    "registration_banner": ["Banner on the event registration page (typically live ~6 months)."],
    "networking_reception": [
        "On-site networking reception with food & beverage.",
        "Company logo on in-person conference signage.",
        "Literature may be available on high-top tables during the reception.",
        "Includes (2) in-person company representative badges for the whole conference."
    ],
    "networking_activity": [
        "Networking activity/excursion.",
        "High-top table at the activity site.",
        "Insert in branded grab-and-go snack bags plus post-conference activity.",
        "Includes (1) in-person company representative badge."
    ],
    "advisory_board": [
        "3-hour advisory board with room, AV, and food & beverage.",
        "Post-event summary and guaranteed attendance."
    ],
    "non_cme_session": [
        "45-minute non-CME/CE symposium in a non-competitive slot.",
        "Full logistics, room and AV support."
    ],
}

# Discount mapping
DISCOUNT_MAP = {
    "none": 1.00,
    "minus_10": 0.90,
    "minus_15": 0.85,
    "minus_20": 0.80,
    "custom": None,
}

# Sarah's information
SARAH = {
    "name": "Sarah Louden",
    "title": "Founder and Executive Director, Total Health Conferencing",
    "email": "sarah@totalhealthconferencing.com",
    "phone": "C: (561) 331-9615  O: (561) 237-2845  F: (561) 771-1748",
}

# Events list
EVENTS = [
    # 2025 Events
    {"meeting_name": "2025 Astera Cancer Care Annual Retreat", "meeting_date_long": "September 27, 2025", "venue": "Ocean Place Resort and Spa", "city_state": "Long Branch, NJ", "default_tier": "standard_1d"},
    {"meeting_name": "2025 Cancer Updates GU and Lung, Memphis, TN", "meeting_date_long": "September 30, 2025", "venue": "Hilton Memphis", "city_state": "Memphis, TN", "default_tier": "standard_1d"},
    {"meeting_name": "2025 West Oncology APP Dinner", "meeting_date_long": "October 9, 2025", "venue": "Hilton Memphis", "city_state": "Memphis, TN", "default_tier": "standard_1d"},
    {"meeting_name": "2025 Empower (Patient Meeting)", "meeting_date_long": "October 11, 2025", "venue": "Farmer's Table", "city_state": "Boca Raton, FL", "default_tier": "standard_1d"},
    {"meeting_name": "2025 Northwell Health Second Annual Liver Cancer Symposium", "meeting_date_long": "October 17, 2025", "venue": "The Garden City Hotel", "city_state": "Garden City, NY", "default_tier": "standard_1d"},
    {"meeting_name": "2025 Cancer Updates GI and Breast, Boston, MA", "meeting_date_long": "October 30, 2025", "venue": "Battery Wharf Hotel Boston Waterfront", "city_state": "Boston, MA", "default_tier": "standard_1d"},
    {"meeting_name": "2025 ESMO USA West", "meeting_date_long": "November 1–2, 2025", "venue": "The Antlers", "city_state": "Colorado Springs, CO", "default_tier": "standard_2d"},
    {"meeting_name": "2025 ESMO USA East", "meeting_date_long": "November 1–2, 2025", "venue": "The Ritz-Carlton Orlando", "city_state": "Orlando, FL", "default_tier": "standard_2d"},
    {"meeting_name": "2025 Northwell Health Multidisciplinary Head and Neck Cancer Symposium", "meeting_date_long": "November 1, 2025", "venue": "The Garden City Hotel", "city_state": "Garden City, NY", "default_tier": "standard_1d"},
    {"meeting_name": "2025 Cancer Updates Heme and GU, Denver, CO", "meeting_date_long": "December 4, 2025", "venue": "Denver Marriott Westminster", "city_state": "Denver, CO", "default_tier": "standard_1d"},
    {"meeting_name": "2025 Northwell Best of Practice Impacting Science", "meeting_date_long": "December 5, 2025", "venue": "The Garden City Hotel", "city_state": "Garden City, NY", "default_tier": "standard_1d"},
    {"meeting_name": "2025 Oncology Update Conference presented by High Plains Oncology Professionals-HPOP", "meeting_date_long": "December 6, 2025", "venue": "Fort Collins Marriott", "city_state": "Fort Collins, CO", "default_tier": "standard_1d"},
    {"meeting_name": "2025 (Northwell) 2nd Annual New York Pancreatic Cancer Consortium", "meeting_date_long": "December 12, 2025", "venue": "Martinique New York on Broadway", "city_state": "Manhattan, NY", "default_tier": "platinum"},
    
    # 2026 Events
    {"meeting_name": "2026 Best of Breast Conference", "meeting_date_long": "January 17-18, 2026", "venue": "Beach House Ft. Lauderdale", "city_state": "Fort Lauderdale, FL", "default_tier": "best_of", "expected_attendance": 50},
    {"meeting_name": "2026 Cancer Updates GI and Breast, Princeton, NJ", "meeting_date_long": "January 22, 2026", "venue": "Princeton Marriott at Forrestal", "city_state": "Princeton, NJ", "default_tier": "standard_1d", "expected_attendance": 30},
    {"meeting_name": "2026 Northwell Cancer Institute Frontiers in Cancer Scientific Symposium", "meeting_date_long": "January 28, 2026", "venue": "Carnegie Hall", "city_state": "Manhattan, NYC", "default_tier": "standard_1d"},
    {"meeting_name": "2026 West Oncology Conference", "meeting_date_long": "January 31-February 1, 2026", "venue": "Hilton Memphis", "city_state": "Memphis, TN", "default_tier": "standard_2d", "expected_attendance": 50},
    {"meeting_name": "2026 Best of Hematology Conference", "meeting_date_long": "February 7-8, 2026", "venue": "The Hythe", "city_state": "Vail, CO", "default_tier": "best_of", "expected_attendance": 30},
    {"meeting_name": "2026 ASCO Direct GI", "meeting_date_long": "February 7-8, 2026", "venue": "The Rittenhouse", "city_state": "Philadelphia, PA", "default_tier": "standard_2d", "expected_attendance": 40},
    {"meeting_name": "2026 University of Kansas Cancer Center Breast Cancer Year in Review", "meeting_date_long": "February 21, 2026", "venue": "Sheraton Overland Park Hotel at the Convention Center", "city_state": "Kansas City, MO", "default_tier": "standard_1d", "expected_attendance": 75},
    {"meeting_name": "2026 Cancer Updates GI and Breast, Dallas, TX", "meeting_date_long": "February 26, 2026", "venue": "The Highland Dallas", "city_state": "Dallas, TX", "default_tier": "standard_1d", "expected_attendance": 30},
    {"meeting_name": "2026 Cancer Updates GI and Breast, Nashville, TN", "meeting_date_long": "March 5, 2026", "venue": "Loews Vanderbilt Hotel", "city_state": "Nashville, TN", "default_tier": "standard_1d", "expected_attendance": 30},
    {"meeting_name": "2026 Cancer Updates GI and Lung, Denver, CO", "meeting_date_long": "March 12, 2026", "venue": "Denver Marriott Westminster", "city_state": "Denver, CO", "default_tier": "standard_1d", "expected_attendance": 30},
    {"meeting_name": "2026 ASCO Direct GU", "meeting_date_long": "March 14-15, 2026", "venue": "The Hyatt", "city_state": "Boston, MA", "default_tier": "standard_2d", "expected_attendance": 40},
    {"meeting_name": "2026 Cancer Updates GI and Breast, Houston", "meeting_date_long": "April 9, 2026", "venue": "The Blossom Hotel Houston", "city_state": "Houston, TX", "default_tier": "standard_1d", "expected_attendance": 30},
    {"meeting_name": "2026 Pathways Conference", "meeting_date_long": "April 10-11, 2026", "venue": "Houston, TX", "city_state": "Houston, TX", "default_tier": "standard_2d"},
    {"meeting_name": "2026 Oncology Clinical Updates - Review and Renew Sedona", "meeting_date_long": "April 11-12, 2026", "venue": "Sedona, AZ", "city_state": "Sedona, AZ", "default_tier": "standard_2d"},
    {"meeting_name": "ESMO in Focus 2026: Lung Cancer and Sarcomas - With John Trent and Ravi Salgia", "meeting_date_long": "April 25, 2026", "venue": "Minneapolis, MN", "city_state": "Minneapolis, MN", "default_tier": "standard_1d"},
    {"meeting_name": "2026 Cancer Updates Breast and Lung, Kansas City", "meeting_date_long": "May 14, 2026", "venue": "Kansas City, MO", "city_state": "Kansas City, MO", "default_tier": "standard_1d"},
    {"meeting_name": "ESMO in Focus 2026: Breast - EAST - With Reshma Mahtani and Kevin Kalinsky", "meeting_date_long": "May 16, 2026", "venue": "Baltimore, MD", "city_state": "Baltimore, MD", "default_tier": "standard_1d"},
    {"meeting_name": "ESMO in Focus 2026: Breast - WEST - With Hope Rugo and Sarah Tolaney", "meeting_date_long": "May 16, 2026", "venue": "Salt Lake City, UT", "city_state": "Salt Lake City, UT", "default_tier": "standard_1d"},
    {"meeting_name": "2026 Best Of ASCO Puerto Rico", "meeting_date_long": "June 13-14, 2026", "venue": "Hyatt Regency Grand Reserve", "city_state": "San Juan, PR", "default_tier": "standard_2d", "expected_attendance": 60},
    {"meeting_name": "2026 Best Of ASCO Austin", "meeting_date_long": "June 13-14, 2026", "venue": "W Austin", "city_state": "Austin, TX", "default_tier": "standard_2d", "expected_attendance": 60},
    {"meeting_name": "2026 Best Of ASCO Los Angeles", "meeting_date_long": "June 13-14, 2026", "venue": "Le Meridien", "city_state": "Los Angeles, CA", "default_tier": "standard_2d", "expected_attendance": 50},
    {"meeting_name": "2026 Best Of ASCO Hawaii", "meeting_date_long": "June 27-28, 2026", "venue": "Sheraton Waikiki Beach Resort", "city_state": "Honolulu, HI", "default_tier": "standard_2d", "expected_attendance": 100},
    {"meeting_name": "2026 Best Of ASCO Washington DC", "meeting_date_long": "June 27-28, 2026", "venue": "Four Seasons Hotel Washington", "city_state": "Washington DC", "default_tier": "standard_2d", "expected_attendance": 60},
    {"meeting_name": "2026 Best Of ASCO Denver", "meeting_date_long": "June 27-28, 2026", "venue": "Denver Marriott Westminster", "city_state": "Denver, CO", "default_tier": "standard_2d", "expected_attendance": 60},
    {"meeting_name": "2026 Cancer Updates Heme and Breast, Michigan", "meeting_date_long": "July 9, 2026", "venue": "Ann Arbor, MI", "city_state": "Ann Arbor, MI", "default_tier": "standard_1d"},
    {"meeting_name": "2026 ASCO Direct New Orleans", "meeting_date_long": "July 10-11, 2026", "venue": "New Orleans, LA", "city_state": "New Orleans, LA", "default_tier": "standard_2d"},
    {"meeting_name": "2026 ASCO Direct Charlotte", "meeting_date_long": "July 18-19, 2026", "venue": "Grand Bohemian", "city_state": "Charlotte, NC", "default_tier": "standard_2d", "expected_attendance": 60},
    {"meeting_name": "2026 Best Of ASCO Philadelphia", "meeting_date_long": "July 18-19, 2026", "venue": "Rittenhouse", "city_state": "Philadelphia, PA", "default_tier": "standard_2d", "expected_attendance": 60},
    {"meeting_name": "2026 Best Of ASCO Anchorage", "meeting_date_long": "July 18-19, 2026", "venue": "Anchorage Marriott Downtown", "city_state": "Anchorage, AK", "default_tier": "standard_2d", "expected_attendance": 30},
    {"meeting_name": "2026 ASCO Direct Las Vegas", "meeting_date_long": "July 25-26, 2026", "venue": "Westin Lake Las Vegas", "city_state": "Henderson/Las Vegas, NV", "default_tier": "standard_2d", "expected_attendance": 60},
    {"meeting_name": "2026 Best Of ASCO Indianapolis", "meeting_date_long": "July 25-26, 2026", "venue": "Indianapolis Marriott Downtown", "city_state": "Indianapolis, IN", "default_tier": "standard_2d", "expected_attendance": 50},
    {"meeting_name": "2026 Best Of ASCO Minneapolis", "meeting_date_long": "August 1-2, 2026", "venue": "Four Seasons Minneapolis", "city_state": "Minneapolis, MN", "default_tier": "standard_2d", "expected_attendance": 50},
    {"meeting_name": "2026 Best Of ASCO Maine", "meeting_date_long": "August 1-2, 2026", "venue": "TBD", "city_state": "Portland, ME", "default_tier": "standard_2d", "expected_attendance": 50},
    {"meeting_name": "2026 ASCO Direct Boston", "meeting_date_long": "August 8-9, 2026", "venue": "Ritz Carlton Boston", "city_state": "Boston, MA", "default_tier": "standard_2d", "expected_attendance": 50},
    {"meeting_name": "2026 ASCO Direct Seattle", "meeting_date_long": "August 8-9, 2026", "venue": "Hilton Motif Seattle", "city_state": "Seattle, WA", "default_tier": "standard_2d", "expected_attendance": 50},
    {"meeting_name": "2026 Best Of ASCO Memphis", "meeting_date_long": "August 8-9, 2026", "venue": "Hilton Memphis", "city_state": "Memphis, TN", "default_tier": "standard_2d", "expected_attendance": 60},
    {"meeting_name": "2026 Cancer Update GU and Lung, Princeton, NJ", "meeting_date_long": "August 13, 2026", "venue": "Princeton, NJ", "city_state": "Princeton, NJ", "default_tier": "standard_1d"},
    {"meeting_name": "2026 Oncology Clinical Updates - Review and Renew Martha's Vineyard", "meeting_date_long": "August 15-16, 2026", "venue": "Martha's Vineyard, MA", "city_state": "Martha's Vineyard, MA", "default_tier": "standard_2d"},
    {"meeting_name": "2026 Cancer Updates GU and Heme, Houston, TX", "meeting_date_long": "September 10, 2026", "venue": "Houston, TX", "city_state": "Houston, TX", "default_tier": "standard_1d"},
    {"meeting_name": "2026 MDONS Conference", "meeting_date_long": "September 18, 2026", "venue": "Denver, CO", "city_state": "Denver, CO", "default_tier": "standard_1d"},
    {"meeting_name": "2026 Cancer Updates Lung and Breast, Dallas, TX", "meeting_date_long": "September 24, 2026", "venue": "Dallas, TX", "city_state": "Dallas, TX", "default_tier": "standard_1d"},
    {"meeting_name": "2026 Cancer Updates GU and Lung, Philadelphia, PA", "meeting_date_long": "September 24, 2026", "venue": "Philadelphia, PA", "city_state": "Philadelphia, PA", "default_tier": "standard_1d"},
    {"meeting_name": "2026 Astera Cancer Care Annual Retreat", "meeting_date_long": "September 26, 2026", "venue": "Long Branch, NJ", "city_state": "Long Branch, NJ", "default_tier": "standard_1d"},
    {"meeting_name": "2026 Cancer Updates GI and Breast, Boston, MA", "meeting_date_long": "October 1, 2026", "venue": "Boston, MA", "city_state": "Boston, MA", "default_tier": "standard_1d"},
    {"meeting_name": "2026 West Oncology APP Dinner", "meeting_date_long": "October 8, 2026", "venue": "Memphis, TN", "city_state": "Memphis, TN", "default_tier": "standard_1d"},
    {"meeting_name": "2026 Virtual West Oncology APP", "meeting_date_long": "October 9, 2026", "venue": "Virtual", "city_state": "Virtual", "default_tier": "standard_1d"},
    {"meeting_name": "2026 Empower (Patient Meeting)", "meeting_date_long": "October 10, 2026", "venue": "Boca Raton, FL", "city_state": "Boca Raton, FL", "default_tier": "standard_1d"},
    {"meeting_name": "2026 Empower GI Conference", "meeting_date_long": "November 7, 2026", "venue": "Washington DC", "city_state": "Washington DC", "default_tier": "standard_1d"},
    {"meeting_name": "2026 ESMO USA West", "meeting_date_long": "November 14-15, 2026", "venue": "Colorado Springs, CO", "city_state": "Colorado Springs, CO", "default_tier": "standard_2d"},
    {"meeting_name": "2026 ESMO USA East", "meeting_date_long": "November 14-15, 2026", "venue": "Orlando, FL", "city_state": "Orlando, FL", "default_tier": "standard_2d"},
    {"meeting_name": "2026 Multidisciplinary Head and Neck Cancer Symposium", "meeting_date_long": "November 14, 2026", "venue": "Garden City, NY", "city_state": "Garden City, NY", "default_tier": "standard_1d"},
]

# Default asset paths
DEFAULT_LOGO_PATHS = ["th_logo.png", "th_logo.jpg", "TH Logo.png", "TH Logo.jpg"]
DEFAULT_SIG_PATHS = ["sarah_signature.png", "sarah_signature.jpg", "Sarah signature.png", "Sarah signature.jpg"]

# ---------------------- Helper Functions ----------------------

def currency(n: float) -> str:
    """Format number as currency"""
    return f"${n:,.2f}"

def round_nearest_50(n: float) -> float:
    """Round to nearest $50"""
    return round(n / 50.0) * 50.0

def booth_bullets() -> List[str]:
    """Standard booth benefits"""
    return [
        "In-person exhibit booth — (1) 6' draped table and 2 chairs.",
        "(2) Full registration admissions for company representatives; additional badges available for purchase.",
        "Company logo on in-person and virtual signage.",
        "Company logo on the conference app.",
        "(1) Conference bag insert.",
        "Pre- and post-conference registration list.",
    ]

def read_first_existing(paths) -> Optional[bytes]:
    """Read first existing file from list of paths"""
    for p in paths:
        fp = Path(p)
        if fp.exists() and fp.is_file():
            return fp.read_bytes()
    return None

# Embedded images for Streamlit Cloud deployment
def get_embedded_logo() -> Optional[bytes]:
    """Get embedded logo for cloud deployment"""
    try:
        # Try to read from file first
        logo_bytes = read_first_existing(DEFAULT_LOGO_PATHS)
        if logo_bytes:
            return logo_bytes
    except:
        pass
    # Fallback: return None (no logo)
    return None

def get_embedded_signature() -> Optional[bytes]:
    """Get embedded signature for cloud deployment"""
    try:
        # Try to read from file first
        sig_bytes = read_first_existing(DEFAULT_SIG_PATHS)
        if sig_bytes:
            return sig_bytes
    except:
        pass
    # Fallback: return None (no signature)
    return None

def render_letter_paragraphs(payload: Dict, document_type: str = 'LOR') -> List[str]:
    """Generate letter paragraphs from payload"""
    if document_type == 'LOA':
        return render_loa_paragraphs(payload)
    else:
        return render_lor_paragraphs(payload)

def render_lor_paragraphs(payload: Dict) -> List[str]:
    """Generate LOR paragraphs from payload"""
    today_long = dt.date.today().strftime("%B %d, %Y")
    audience = payload.get(
        "audience_list",
        "physicians, nurses, pharmacists, advanced practitioners and patient advocates",
    )
    attendance_expected = payload.get("attendance_expected")

    paras = []
    paras.append("Letter of Request")
    paras.append(today_long)
    paras.append("Dear Exhibitor,")
    paras.append(
        f"Total Health Conferencing is proud to submit this request to {payload['company_name']} for support of {payload['meeting_name']}. "
        f"The meeting will take place on {payload['meeting_date_long']} at {payload['venue']}, {payload['city_state']}."
    )
    if attendance_expected:
        paras.append(
            f"We expect {attendance_expected} total attendees, including {audience}. Attendance numbers are expected, but not guaranteed."
        )
    else:
        paras.append(
            f"We expect a strong mix of {audience}. Attendance numbers are expected, but not guaranteed."
        )
    paras.append(
        f"{payload['meeting_name']} will showcase clinical presentations, allowing attendees to engage in presentation, discussion, analysis, and participation."
    )
    paras.append(
        f"Your support in the amount of {payload['amount_currency']} will provide you with an opportunity to support high-quality education. "
        f"Total Health Conferencing will provide you with the following benefits:"
    )
    return paras

# ---------------------- LOA Generation Functions ----------------------

def render_loa_paragraphs(payload: Dict, booth_selected: bool = False, add_on_keys: List[str] = None, current_add_ons: Dict = None) -> List[str]:
    """Generate LOA paragraphs based on the template structure"""
    agreement_date = payload.get("agreement_date", dt.date.today().strftime("%B %d, %Y"))
    company_name = payload.get("company_name", "[Company]")
    company_address = payload.get("company_address", "[Address]")
    meeting_name = payload.get("meeting_name", "[Meeting]")
    meeting_date = payload.get("meeting_date_long", "[Date]")
    venue = payload.get("venue", "[Venue]")
    city_state = payload.get("city_state", "[City, State]")
    amount = payload.get("amount_currency", "[Amount]")
    
    # Smart venue display for LOA
    if venue == city_state or venue in city_state:
        location = city_state
    else:
        location = f"{venue}, {city_state}"
    
    paras = []
    paras.append("LETTER OF AGREEMENT (LOA)")
    paras.append("")
    paras.append(f"This Letter of Agreement (the \"Agreement\") is made as of {agreement_date} by and between Total Health Information Services, LLC. (\"Total Health\"), with its principal place of business at 20423 State Road 7, F6-496, Boca Raton FL 33498, and {company_name}(\"Sponsor\"), with its principal place of business at {company_address}. The individual signing this Agreement represents that they have the authority to legally bind the Sponsor to this Agreement.")
    paras.append("")
    paras.append("1. Purpose")
    paras.append("")
    paras.append("The purpose of this Agreement is to outline the terms under which the Sponsor agrees to participate in Total Health's educational events by securing exhibit tables, product theaters, and other sponsorship-related opportunities, as defined in the attached Scope of Work (SOW).")
    paras.append("")
    paras.append("2. Scope of Work (SOW)")
    paras.append("")
    paras.append("The SOW, attached as Exhibit A, details the services Total Health will provide educational related services.")
    paras.append("")
    paras.append(f"Each sponsorship opportunity including name and date of meeting, city, and specific sponsor items will be specified in the SOW. This Agreement applies to all events listed within the specified dates and covers all agreed-upon sponsorship activities.")
    paras.append("")
    paras.append("**SPECIFIC EVENT DETAILS:**")
    paras.append(f"• Event: {meeting_name}")
    paras.append(f"• Date: {meeting_date}")
    paras.append(f"• Location: {location}")
    paras.append(f"• Total Sponsorship Amount: {amount}")
    paras.append("")
    
    # Add detailed SOW section with all sponsorship components
    paras.append("**DETAILED SCOPE OF WORK:**")
    paras.append("")
    
    # Booth details
    if booth_selected:
        paras.append("**EXHIBIT BOOTH SPONSORSHIP:**")
        booth_tier = payload.get("booth_tier", "Standard")
        booth_price = payload.get("booth_price", "TBD")
        paras.append(f"• Booth Tier: {booth_tier}")
        paras.append(f"• Booth Cost: {booth_price}")
        paras.append("• Booth Benefits:")
        paras.append("  - Designated exhibit space at the event")
        paras.append("  - Company name and logo in event materials")
        paras.append("  - Access to attendee networking opportunities")
        paras.append("  - Inclusion in event directory and program")
        paras.append("")
    
    # Add-ons details
    if add_on_keys and current_add_ons:
        paras.append("**ADDITIONAL SPONSORSHIP COMPONENTS:**")
        for k in add_on_keys:
            addon = current_add_ons[k]
            paras.append(f"• {addon['label']}: {currency(addon['price'])}")
            # Add specific benefits for each add-on
            if k == "program_ad_full":
                paras.append("  - Full-page advertisement in printed/digital program guide")
            elif k == "charging_stations":
                paras.append("  - Branded charging station with company artwork")
                paras.append("  - Includes (1) in-person company representative badge")
            elif k == "wifi_sponsorship":
                paras.append("  - Exclusive Wi-Fi sponsorship with company logo/name")
                paras.append("  - Includes (1) in-person company representative badge")
            elif k == "platform_banner":
                paras.append("  - Banner advertisement on event's digital platform/lobby page")
            elif k == "email_banner":
                paras.append("  - Banner placement in national call-to-action email")
            elif k == "registration_banner":
                paras.append("  - Banner on event registration page (typically live ~6 months)")
            elif k == "networking_reception":
                paras.append("  - On-site networking reception with food & beverage")
                paras.append("  - Company logo on in-person conference signage")
                paras.append("  - Literature may be available on high-top tables")
                paras.append("  - Includes (2) in-person company representative badges")
            elif k == "networking_activity":
                paras.append("  - Networking activity/excursion")
                paras.append("  - High-top table at the activity site")
                paras.append("  - Insert in branded grab-and-go snack bags")
                paras.append("  - Includes (1) in-person company representative badge")
            elif k == "advisory_board":
                paras.append("  - 3-hour advisory board with room, AV, and food & beverage")
                paras.append("  - Post-event summary and guaranteed attendance")
            elif k == "non_cme_session":
                paras.append("  - 45-minute non-CME/CE symposium in non-competitive slot")
                paras.append("  - Full logistics, room and AV support")
        paras.append("")
    
    # Additional information if provided
    if payload.get("additional_info"):
        paras.append("**ADDITIONAL SPONSORSHIP REQUIREMENTS:**")
        paras.append(payload["additional_info"])
        paras.append("")
    paras.append("3. Payment and Financial Terms")
    paras.append("")
    paras.append("Full payment is due within 60 days of receiving the invoice, regardless of the Sponsor's attendance or participation. No refunds shall be issued under any circumstances, except those agreed upon in writing.")
    paras.append("")
    paras.append("Any notice of cancellation for participation in an event must be submitted in writing at least sixty (60) days prior to the event. Cancellations made within this period will receive a 50% credit of all fees paid toward a future program, at Total Health's discretion. Cancellations after this period will not be entitled to any credit. In cases where Total Health is unable to produce a scheduled meeting, a credit may be issued at Total Health's toward future programs.")
    paras.append("")
    paras.append("4. Right to Refuse Exhibit")
    paras.append("")
    paras.append("Total Health reserves the right to decline, prohibit, or expel any exhibit it deems inappropriate or out of character with the event, or if the exhibit violates the terms of this Agreement or applicable laws and regulations. No refunds or credits will be issued if an exhibit is refused or expelled under these conditions.")
    paras.append("")
    paras.append("5. Compliance with Laws and Facility Rules")
    paras.append("")
    paras.append("The Sponsor shall comply with all applicable laws, codes, and regulations, as well as the rules of the venue where the event is held. The Sponsor assumes all liability for any non-compliance and agrees to indemnify Total Health against any claims arising from such violations.")
    paras.append("")
    paras.append("6. Advertising and Solicitation")
    paras.append("")
    paras.append("Distribution of advertising materials and solicitation is restricted to the Sponsor's designated booth area only. Any unauthorized promotion outside the assigned space may result in the removal of the Sponsor from the event, without refund or credit.")
    paras.append("")
    paras.append("7. Occupancy of Exhibit Space")
    paras.append("")
    paras.append("Sponsor's use of exhibit space is mandatory. Should the Sponsor fail to occupy the space, Total Health reserves the right to repurpose the space as it sees fit, without offering any rebate or credit.")
    paras.append("")
    paras.append("8. Installation and Dismantling of Exhibits")
    paras.append("")
    paras.append("The Sponsor is responsible for adhering to the installation and dismantling schedules provided by Total Health. Any deviations or delays in setup or breakdown may result in penalties or additional charges. It is the Sponsor's responsibility to coordinate with their team and notify any third parties involved about these schedules.")
    paras.append("")
    paras.append("9. Liability and Insurance")
    paras.append("")
    paras.append("Total Health and the venue assume no responsibility for the protection or safety of the Sponsor's representatives, agents, or property. Total Health recommends that small and easily portable items be secured or removed outside of exhibit hours.")
    paras.append("")
    paras.append("Indemnification: The Sponsor agrees to indemnify, defend, and hold harmless Total Health, its officers, employees, and agents from any claims, damages, or liabilities arising from the Sponsor's participation, exhibit, or actions at the event.")
    paras.append("")
    paras.append("11. Force Majeure")
    paras.append("")
    paras.append("Total Health shall not be liable for any failure to perform its obligations under this Agreement due to circumstances beyond its reasonable control, including, but not limited to, acts of God, natural disasters, pandemics, government restrictions, or any other unforeseen event (\"Force Majeure\"). In such cases, Total Health may reschedule the event or offer a credit at its discretion, without liability for any resulting damages or costs incurred by the Sponsor.")
    paras.append("")
    paras.append("12. Confidentiality")
    paras.append("")
    paras.append("The terms of this Agreement, including the SOW and all proprietary information exchanged between Total Health and the Sponsor, shall be considered confidential and shall not be disclosed to third parties without the prior written consent of both parties, except as required by law.")
    paras.append("")
    paras.append("13. Entire Agreement")
    paras.append("")
    paras.append("This Agreement, together with the attached SOW, represents the complete understanding between Total Health and the Sponsor, superseding any prior discussions or agreements. Any amendments must be in writing and signed by both parties to be valid.")
    paras.append("")
    paras.append("14. Governing Law and Dispute Resolution")
    paras.append("")
    paras.append("This Agreement shall be governed by the laws of the State of Florida. Any disputes arising out of or relating to this Agreement shall be resolved through binding arbitration in West Palm Beach, FL, with each party bearing its own legal fees and costs.")
    paras.append("")
    paras.append("15. Signatures")
    paras.append("")
    paras.append("By signing below, the parties agree to the terms and conditions outlined in this Agreement.")
    paras.append("")
    paras.append("Total Health")
    paras.append("")
    paras.append("By: ____________________________")
    # Split name and title properly
    signature_person = payload.get('signature_person', 'Sarah Louden - Founder and Executive Director, Total Health Conferencing')
    if ' - ' in signature_person:
        name, title = signature_person.split(' - ', 1)
        paras.append(f"Name: {name}")
        paras.append(f"Title: {title}")
    else:
        paras.append(f"Name: {signature_person}")
        paras.append("Title: ____________________________")
    paras.append(f"Date: {dt.date.today().strftime('%B %d, %Y')}")
    paras.append("")
    paras.append("")
    paras.append("")
    paras.append("")
    paras.append(f"{company_name}")
    paras.append("")
    paras.append("By: ____________________________")
    paras.append("Name: ____________________________")
    paras.append("Title: ____________________________")
    paras.append("Date: ___________________________")
    
    return paras

def build_loa_docx_bytes(paragraphs: List[str], logo_bytes: Optional[bytes], sig_bytes: Optional[bytes] = None) -> bytes:
    """Generate LOA DOCX document"""
    if not DOCX_AVAILABLE:
        return "\n".join(paragraphs).encode("utf-8")

    doc = Document()

    # Logo (centered, proper aspect ratio)
    if logo_bytes:
        try:
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(BytesIO(logo_bytes), width=Inches(2.5))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
        except Exception:
            pass

    # Add paragraphs
    for para_text in paragraphs:
        if para_text.strip():
            if para_text.startswith("LETTER OF AGREEMENT"):
                # Title - centered and bold
                p = doc.add_paragraph(para_text)
                p.runs[0].bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(12)
            elif para_text.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "11.", "12.", "13.", "14.", "15.")):
                # Section headers - bold headings
                p = doc.add_paragraph(para_text)
                p.runs[0].bold = True
                p.paragraph_format.space_after = Pt(6)
            elif para_text.startswith("**SPECIFIC EVENT DETAILS:**"):
                # Event details header
                p = doc.add_paragraph("SPECIFIC EVENT DETAILS:")
                p.runs[0].bold = True
                p.paragraph_format.space_after = Pt(6)
            elif para_text.startswith("•"):
                # Bullet points
                p = doc.add_paragraph(para_text)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(3)
            elif para_text.startswith(("Total Health", "By:", "Name:", "Title:", "Date:")):
                # Signature section - field labels
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(6)
            elif para_text.strip() and not para_text.startswith(("By:", "Name:", "Title:", "Date:", "•", "**", "1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "11.", "12.", "13.", "14.", "15.", "LETTER OF AGREEMENT", "SPECIFIC EVENT DETAILS", "DETAILED SCOPE OF WORK", "Total Health")) and len(para_text.strip()) < 50 and not " " in para_text.strip():
                # Company names in signature section (sponsor company name) - short text without spaces
                p = doc.add_paragraph(para_text)
                p.runs[0].bold = True
                p.paragraph_format.space_after = Pt(6)
            else:
                # Regular paragraphs
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(6)

    # No signature image for LOA - signatures are filled in manually

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

def build_loa_pdf_bytes(paragraphs: List[str], logo_bytes: Optional[bytes], sig_bytes: Optional[bytes] = None) -> bytes:
    """Generate LOA PDF document"""
    if not PDF_AVAILABLE:
        return "\n".join(paragraphs).encode("utf-8")

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    story = []

    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=16, spaceAfter=12, alignment=TA_CENTER)
    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=12, spaceAfter=6, alignment=TA_LEFT)
    body_style = ParagraphStyle('CustomBody', parent=styles['Normal'], fontSize=10, spaceAfter=6, alignment=TA_LEFT)
    bullet_style = ParagraphStyle('CustomBullet', parent=styles['Normal'], fontSize=10, spaceAfter=3, alignment=TA_LEFT, leftIndent=20)

    # Logo (centered, proper aspect ratio like LOR)
    if logo_bytes and PIL_AVAILABLE:
        try:
            img = Image.open(BytesIO(logo_bytes))
            w, h = img.size
            max_w = 2.5 * inch
            ratio = max_w / w
            story.append(RLImage(BytesIO(logo_bytes), width=max_w, height=h * ratio, hAlign="CENTER"))
            story.append(Spacer(1, 0.20 * inch))
        except Exception:
            pass

    # Add paragraphs
    for para_text in paragraphs:
        if para_text.strip():
            if para_text.startswith("LETTER OF AGREEMENT"):
                story.append(Paragraph(para_text, title_style))
            elif para_text.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "11.", "12.", "13.", "14.", "15.")):
                story.append(Paragraph(para_text, heading_style))
            elif para_text.startswith("**SPECIFIC EVENT DETAILS:**"):
                story.append(Paragraph("SPECIFIC EVENT DETAILS:", heading_style))
            elif para_text.startswith("•"):
                story.append(Paragraph(para_text, bullet_style))
            elif para_text.startswith(("Total Health", "By:", "Name:", "Title:", "Date:")):
                # Signature section - field labels
                story.append(Paragraph(para_text, body_style))
            elif para_text.strip() and not para_text.startswith(("By:", "Name:", "Title:", "Date:", "•", "**", "1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "11.", "12.", "13.", "14.", "15.", "LETTER OF AGREEMENT", "SPECIFIC EVENT DETAILS", "DETAILED SCOPE OF WORK", "Total Health")) and len(para_text.strip()) < 50 and not " " in para_text.strip():
                # Company names in signature section (sponsor company name) - short text without spaces
                story.append(Paragraph(para_text, heading_style))
            else:
                story.append(Paragraph(para_text, body_style))

    # No signature image for LOA - signatures are filled in manually

    doc.build(story)
    return buf.getvalue()

def parse_additional_info(text: str):
    """Parse additional info text into lead-in and bullets"""
    if not text:
        return None, []
    lines = [ln.strip() for ln in text.splitlines()]
    lines = [ln for ln in lines if ln]
    if not lines:
        return None, []
    lead_in = lines[0]
    cleaned = []
    for bl in lines[1:]:
        s = bl.lstrip("-•*").strip()
        if len(s) > 2 and s[0].isdigit() and s[1] in (".", ")"):
            s = s[2:].strip()
        cleaned.append(s)
    return lead_in, cleaned

def get_event_year(event_name: str) -> int:
    """Extract year from event name"""
    if "2026" in event_name:
        return 2026
    elif "2025" in event_name:
        return 2025
    else:
        return 2025  # Default to 2025

def parse_event_date(date_str: str) -> datetime:
    """Parse event date string into datetime object"""
    try:
        # Handle different date formats
        if "–" in date_str or "-" in date_str:
            # Multi-day events - use the first date
            date_part = date_str.split("–")[0].split("-")[0].strip()
        else:
            date_part = date_str.strip()
        
        # Parse the date
        return datetime.strptime(date_part, "%B %d, %Y")
    except:
        # Fallback - try to extract year and month
        try:
            if "January" in date_str:
                return datetime(2025, 1, 1) if "2025" in date_str else datetime(2026, 1, 1)
            elif "February" in date_str:
                return datetime(2025, 2, 1) if "2025" in date_str else datetime(2026, 2, 1)
            elif "March" in date_str:
                return datetime(2025, 3, 1) if "2025" in date_str else datetime(2026, 3, 1)
            elif "April" in date_str:
                return datetime(2025, 4, 1) if "2025" in date_str else datetime(2026, 4, 1)
            elif "May" in date_str:
                return datetime(2025, 5, 1) if "2025" in date_str else datetime(2026, 5, 1)
            elif "June" in date_str:
                return datetime(2025, 6, 1) if "2025" in date_str else datetime(2026, 6, 1)
            elif "July" in date_str:
                return datetime(2025, 7, 1) if "2025" in date_str else datetime(2026, 7, 1)
            elif "August" in date_str:
                return datetime(2025, 8, 1) if "2025" in date_str else datetime(2026, 8, 1)
            elif "September" in date_str:
                return datetime(2025, 9, 1) if "2025" in date_str else datetime(2026, 9, 1)
            elif "October" in date_str:
                return datetime(2025, 10, 1) if "2025" in date_str else datetime(2026, 10, 1)
            elif "November" in date_str:
                return datetime(2025, 11, 1) if "2025" in date_str else datetime(2026, 11, 1)
            elif "December" in date_str:
                return datetime(2025, 12, 1) if "2025" in date_str else datetime(2026, 12, 1)
            else:
                return datetime.now()
        except:
            return datetime.now()

def get_upcoming_events() -> List[Dict]:
    """Get all events, sorted chronologically (includes past events)"""
    all_events = []
    
    for event in EVENTS:
        all_events.append(event)
    
    # Sort by date
    all_events.sort(key=lambda x: parse_event_date(x["meeting_date_long"]))
    return all_events

# ---------------------- DOCX Export Functions ----------------------

def _docx_add_bullets(doc, items: List[str]):
    """Add bullet points to DOCX document"""
    for text in items:
        p = doc.add_paragraph(text, style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.30)
        p.paragraph_format.first_line_indent = Inches(-0.20)
        p.paragraph_format.space_after = Pt(2)

def _docx_add_header(doc, title: str):
    """Add header to DOCX document"""
    p = doc.add_paragraph(title)
    if p.runs:
        p.runs[0].bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)

def build_docx_bytes(paragraphs: List[str], booth_selected: bool,
                     add_on_keys: List[str],
                     logo_bytes: Optional[bytes], sig_bytes: Optional[bytes],
                     addl_info: Optional[str]) -> bytes:
    """Build DOCX document from components"""
    if not DOCX_AVAILABLE:
        return "\n".join(paragraphs).encode("utf-8")

    doc = Document()

    # Logo (centered)
    if logo_bytes:
        try:
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(BytesIO(logo_bytes), width=Inches(2.5))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
        except Exception:
            pass

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Body paragraphs
    for i, para in enumerate(paragraphs):
        p = doc.add_paragraph(para)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(8 if i == 0 else 10)

    # Booth
    if booth_selected:
        doc.add_paragraph("").paragraph_format.space_after = Pt(6)
        _docx_add_header(doc, "Exhibit Booth")
        _docx_add_bullets(doc, booth_bullets())

    # Add-ons
    if add_on_keys:
        doc.add_paragraph("").paragraph_format.space_after = Pt(6)
        _docx_add_header(doc, "Selected Add-Ons")
        for k in add_on_keys:
            _docx_add_header(doc, ADD_ONS_2025[k]["label"])  # Use 2025 labels for consistency
            _docx_add_bullets(doc, ADD_ON_BULLETS.get(k, []))

    # Additional Information
    if addl_info:
        doc.add_paragraph("").paragraph_format.space_after = Pt(6)
        _docx_add_header(doc, "Additional Information")
        lead_in, bullets = parse_additional_info(addl_info)
        if lead_in:
            p = doc.add_paragraph(lead_in)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(6)
        if bullets:
            _docx_add_bullets(doc, bullets)
            doc.add_paragraph("").paragraph_format.space_after = Pt(8)

    # Closing + Signature
    doc.add_paragraph("").paragraph_format.space_after = Pt(6)
    closing = doc.add_paragraph("Grateful for your support,")
    closing.alignment = WD_ALIGN_PARAGRAPH.LEFT
    closing.paragraph_format.space_after = Pt(6)

    if sig_bytes:
        try:
            sp = doc.add_paragraph()
            run = sp.add_run()
            run.add_picture(BytesIO(sig_bytes), width=Inches(2.0))
            sp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sp.paragraph_format.space_after = Pt(4)
        except Exception:
            pass

    doc.add_paragraph(SARAH["name"]).alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(SARAH["title"]).alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(SARAH["email"]).alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(SARAH["phone"]).alignment = WD_ALIGN_PARAGRAPH.LEFT

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ---------------------- PDF Export Functions ----------------------

def _pdf_bullet_table(items: List[str], para_style, col_pad_pt: int = 6):
    """Render bullet list as 2-column table: [•] [wrapped text]"""
    if not items:
        return None
    rows = []
    bullet_para = Paragraph("•", para_style)
    for txt in items:
        rows.append([bullet_para, Paragraph(txt, para_style)])
    tbl = Table(rows, colWidths=[12, None])  # 12pt for dot; rest auto
    tbl.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (1, 0), (1, -1), col_pad_pt),
        ("LINEBELOW", (0, 0), (-1, -1), 0, colors.white),
    ]))
    return tbl

def build_pdf_bytes(paragraphs: List[str], booth_selected: bool,
                    add_on_keys: List[str],
                    logo_bytes: Optional[bytes], sig_bytes: Optional[bytes],
                    addl_info: Optional[str]) -> bytes:
    """Build PDF document from components"""
    if not PDF_AVAILABLE:
        return "\n".join(paragraphs).encode("utf-8")

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch
    )

    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["Normal"], fontName="Times-Roman",
                          fontSize=11, leading=16, alignment=TA_LEFT)
    bold = ParagraphStyle("Bold", parent=styles["Normal"], fontName="Times-Bold",
                          fontSize=11, leading=16, alignment=TA_LEFT)

    story = []

    # Logo (centered)
    if logo_bytes and PIL_AVAILABLE:
        try:
            img = Image.open(BytesIO(logo_bytes))
            w, h = img.size
            max_w = 2.5 * inch
            ratio = max_w / w
            story.append(RLImage(BytesIO(logo_bytes), width=max_w, height=h * ratio, hAlign="CENTER"))
            story.append(Spacer(1, 0.20 * inch))
        except Exception:
            pass

    # Body paragraphs
    for i, para in enumerate(paragraphs):
        story.append(Paragraph(para, body))
        story.append(Spacer(1, 0.12 * inch if i != 0 else 0.10 * inch))

    # Booth bullets (table)
    if booth_selected:
        story.append(Spacer(1, 0.10 * inch))
        story.append(Paragraph("Exhibit Booth", bold))
        story.append(Spacer(1, 0.05 * inch))
        tbl = _pdf_bullet_table(booth_bullets(), body)
        if tbl:
            story.append(tbl)

    # Add-on bullets (table per add-on)
    if add_on_keys:
        story.append(Spacer(1, 0.12 * inch))
        story.append(Paragraph("Selected Add-Ons", bold))
        for k in add_on_keys:
            story.append(Spacer(1, 0.08 * inch))
            story.append(Paragraph(ADD_ONS_2025[k]["label"], bold))  # Use 2025 labels
            tbl = _pdf_bullet_table(ADD_ON_BULLETS.get(k, []), body)
            if tbl:
                story.append(tbl)

    # Additional Information (lead-in + table bullets)
    if addl_info:
        story.append(Spacer(1, 0.12 * inch))
        story.append(Paragraph("Additional Information", bold))
        story.append(Spacer(1, 0.04 * inch))
        lead_in, bullets = parse_additional_info(addl_info)
        if lead_in:
            story.append(Paragraph(lead_in, body))
            story.append(Spacer(1, 0.06 * inch))
        if bullets:
            tbl = _pdf_bullet_table(bullets, body)
            if tbl:
                story.append(tbl)
                story.append(Spacer(1, 0.08 * inch))

    # Closing + signature
    story.append(Spacer(1, 0.12 * inch))
    story.append(Paragraph("Grateful for your support,", body))
    story.append(Spacer(1, 0.08 * inch))
    if sig_bytes and PIL_AVAILABLE:
        try:
            img = Image.open(BytesIO(sig_bytes))
            w, h = img.size
            max_w = 2.0 * inch
            ratio = max_w / w
            story.append(RLImage(BytesIO(sig_bytes), width=max_w, height=h * ratio, hAlign="LEFT"))
            story.append(Spacer(1, 0.05 * inch))
        except Exception:
            pass

    for line in (SARAH["name"], SARAH["title"], SARAH["email"], SARAH["phone"]):
        story.append(Paragraph(line, body))

    doc.build(story)
    return buf.getvalue()

# ---------------------- Streamlit UI ----------------------

# Page configuration
st.set_page_config(
    page_title="LOR/LOA Generator v46.0",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Total Health Brand Colors
TH_COLORS = {
    "primary": "#013955",      # Deep teal
    "secondary": "#6d5065",    # Plum
    "accent": "#bb9cb2",       # Lavender
    "blue_gray": "#4a8499",    # Blue-gray
    "light_blue": "#c9deff",   # Light blue
    "light_gray": "#eeedeb",   # Light gray
    "white": "#ffffff",        # White
    "black": "#000000",        # Black
    "dark_gray": "#1a1a1a"     # Dark gray
}

# Custom CSS for Total Health branding
st.markdown(f"""
<style>
    /* Total Health Brand Styling */
    .main {{
        background: #f8f9fa !important;
        min-height: 100vh;
    }}
    
    .main .block-container {{
        background: #f8f9fa !important;
        padding: 2rem;
        max-width: 1200px;
    }}
    
    /* Header Styling */
    .main-header {{
        background: white !important;
        border-radius: 15px;
        padding: 2rem;
        margin-bottom: 2rem;
        text-align: center;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.1);
        border: 2px solid {TH_COLORS['primary']};
    }}
    
    .logo-container {{
        margin-bottom: 1rem;
    }}
    
    .logo-container img {{
        max-width: 480px;
        height: auto;
        display: block;
        margin: 0 auto;
    }}
    
    .logo-container {{
        margin-bottom: 1rem;
        text-align: center;
    }}
    
    h1 {{
        color: {TH_COLORS['primary']};
        font-size: 2.5rem;
        font-weight: 700;
        margin: 0.5rem 0;
        text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.1);
    }}
    
    .version-info {{
        color: {TH_COLORS['secondary']};
        font-size: 0.9rem;
        font-style: italic;
        margin-top: 0.5rem;
    }}
    
    .status-badge {{
        display: inline-block;
        background: linear-gradient(135deg, #28a745, #20c997);
        color: white;
        padding: 0.5rem 1rem;
        border-radius: 25px;
        font-size: 0.9rem;
        font-weight: 600;
        margin-top: 1rem;
        box-shadow: 0 4px 15px rgba(40, 167, 69, 0.3);
    }}
    
    .last-updated {{
        color: {TH_COLORS['secondary']};
        font-size: 0.8rem;
        margin-top: 0.5rem;
    }}
    
    /* Form Styling */
    .stContainer {{
        background: white !important;
        border-radius: 15px;
        padding: 1.5rem;
        margin-bottom: 1.5rem;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.1);
        border: 1px solid {TH_COLORS['light_gray']};
    }}
    
    .section-header {{
        color: {TH_COLORS['primary']};
        font-size: 1.3rem;
        font-weight: 600;
        margin-bottom: 1rem;
        border-bottom: 2px solid {TH_COLORS['accent']};
        padding-bottom: 0.5rem;
    }}
    
    /* Form Controls */
    .stSelectbox > div > div {{
        background: white !important;
        border: 2px solid {TH_COLORS['light_gray']} !important;
        border-radius: 8px !important;
    }}
    
    .stSelectbox > div > div:focus {{
        border-color: {TH_COLORS['primary']} !important;
        box-shadow: 0 0 0 3px rgba(1, 57, 85, 0.1) !important;
    }}
    
    .stTextInput > div > div > input {{
        background: white !important;
        border: 2px solid {TH_COLORS['light_gray']} !important;
        border-radius: 8px !important;
        color: {TH_COLORS['black']} !important;
    }}
    
    .stTextInput > div > div > input:focus {{
        border-color: {TH_COLORS['primary']} !important;
        box-shadow: 0 0 0 3px rgba(1, 57, 85, 0.1) !important;
    }}
    
    .stNumberInput > div > div > input {{
        background: white !important;
        border: 2px solid {TH_COLORS['light_gray']} !important;
        border-radius: 8px !important;
        color: {TH_COLORS['black']} !important;
    }}
    
    .stTextArea > div > div > textarea {{
        background: white !important;
        border: 2px solid {TH_COLORS['light_gray']} !important;
        border-radius: 8px !important;
        color: {TH_COLORS['black']} !important;
    }}
    
    .stMultiselect > div > div {{
        background: white !important;
        border: 2px solid {TH_COLORS['light_gray']} !important;
        border-radius: 8px !important;
    }}
    
    .stRadio > div {{
        background: white !important;
        border: 2px solid {TH_COLORS['light_gray']} !important;
        border-radius: 8px !important;
        padding: 1rem !important;
    }}
    
    /* Button Styling */
    .stButton > button {{
        background: {TH_COLORS['primary']} !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 0.75rem 2rem !important;
        font-weight: 600 !important;
        transition: all 0.3s ease !important;
    }}
    
    .stButton > button:hover {{
        background: {TH_COLORS['secondary']} !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.2) !important;
    }}
    
    /* NUCLEAR READABILITY FIX - FORCE ALL TEXT TO BE READABLE */
    
    /* Force ALL elements to have black text */
    *, *::before, *::after {{
        color: #000000 !important;
    }}
    
    /* Force ALL text elements */
    p, div, span, label, input, textarea, select, option, button, h1, h2, h3, h4, h5, h6,
    .stMarkdown, .stMarkdown p, .stMarkdown div, .stMarkdown span, .stMarkdown strong, .stMarkdown em,
    .stSelectbox, .stTextInput, .stNumberInput, .stTextArea, .stMultiselect, .stRadio, .stButton {{
        color: #000000 !important;
    }}
    
    /* Force ALL form elements to have white backgrounds */
    .stSelectbox, .stTextInput, .stNumberInput, .stTextArea, .stMultiselect, .stRadio, .stButton,
    .stSelectbox > div, .stTextInput > div, .stNumberInput > div, .stTextArea > div, .stMultiselect > div, .stRadio > div {{
        background: white !important;
        color: #000000 !important;
    }}
    
    /* Force ALL inputs and dropdowns */
    input, textarea, select, [role="combobox"], [role="listbox"], [role="option"] {{
        background: white !important;
        color: #000000 !important;
        border: 1px solid #cccccc !important;
    }}
    
    /* Force ALL labels */
    label, .stSelectbox label, .stTextInput label, .stNumberInput label, .stTextArea label, .stMultiselect label, .stRadio label {{
        color: #000000 !important;
        background: transparent !important;
    }}
    
    /* Force ALL success/info messages */
    .stSuccess, .stInfo, .stWarning, .stError, .stAlert {{
        background: white !important;
        color: #000000 !important;
        border: 1px solid #cccccc !important;
    }}
    
    /* Force ALL Streamlit widgets */
    [data-testid="stSelectbox"], [data-testid="stTextInput"], [data-testid="stNumberInput"], 
    [data-testid="stTextArea"], [data-testid="stMultiselect"], [data-testid="stRadio"], 
    [data-testid="stButton"], [data-testid="stMarkdown"] {{
        background: white !important;
        color: #000000 !important;
    }}
    
    /* Force ALL dropdown options */
    [role="menu"], [role="option"], [data-baseweb="select"], [data-baseweb="popover"] {{
        background: white !important;
        color: #000000 !important;
    }}
    
    /* Override any Streamlit default styling */
    .stApp, .main, .block-container {{
        background: #f8f9fa !important;
    }}
    
    /* Force all containers to be white */
    .stContainer, .stSelectbox, .stTextInput, .stNumberInput, .stTextArea, .stMultiselect, .stRadio {{
        background: white !important;
        border: 1px solid #e0e0e0 !important;
    }}
    
    /* Fix button styling - remove white backgrounds from buttons */
    .stButton > button {{
        background: {TH_COLORS['primary']} !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 0.75rem 2rem !important;
        font-weight: 600 !important;
        transition: all 0.3s ease !important;
    }}
    
    .stButton > button:hover {{
        background: {TH_COLORS['secondary']} !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.2) !important;
    }}
    
    /* Remove white backgrounds from success/info messages */
    .stSuccess {{
        background: #d4edda !important;
        color: #155724 !important;
        border: 1px solid #c3e6cb !important;
    }}
    
    .stInfo {{
        background: #d1ecf1 !important;
        color: #0c5460 !important;
        border: 1px solid #bee5eb !important;
    }}
    
    .stWarning {{
        background: #fff3cd !important;
        color: #856404 !important;
        border: 1px solid #ffeaa7 !important;
    }}
    
    .stError {{
        background: #f8d7da !important;
        color: #721c24 !important;
        border: 1px solid #f5c6cb !important;
    }}
    
    /* Fix text color for buttons and other elements */
    .stButton > button, .stButton > button * {{
        color: white !important;
    }}
    
    /* Ensure button text is white */
    button, button * {{
        color: white !important;
    }}
    
    /* Fix any other white text issues */
    .stButton button span {{
        color: white !important;
    }}
    
    /* Fix expandable section headers */
    .streamlit-expanderHeader {{
        color: #000000 !important;
        background: white !important;
    }}
    
    .streamlit-expanderHeader:hover {{
        color: #000000 !important;
        background: #f8f9fa !important;
    }}
    
    /* Fix all expandable content */
    .streamlit-expander {{
        background: white !important;
        color: #000000 !important;
    }}
    
    .streamlit-expanderContent {{
        background: white !important;
        color: #000000 !important;
    }}
    
    /* Fix checkbox labels */
    .stCheckbox label {{
        color: #000000 !important;
        background: transparent !important;
    }}
    
    /* Fix any other dark text issues */
    .stExpander, .stExpander * {{
        color: #000000 !important;
    }}
    
    /* Hide Streamlit branding */
    #MainMenu {{visibility: hidden;}}
    footer {{visibility: hidden;}}
    header {{visibility: hidden;}}
    
    /* Excel Bulk Section Styling */
    .excel-bulk-section {{
        background: linear-gradient(135deg, {TH_COLORS['primary']}15, {TH_COLORS['secondary']}15) !important;
        padding: 1.5rem !important;
        border-radius: 12px !important;
        border: 2px solid {TH_COLORS['light_gray']} !important;
        margin: 1rem 0 !important;
    }}
    
    /* Override all dark elements in Excel bulk section */
    .excel-bulk-section * {{
        color: {TH_COLORS['primary']} !important;
        background-color: transparent !important;
    }}
    
    .excel-bulk-section .stTextInput input {{
        background: white !important;
        border: 2px solid {TH_COLORS['light_gray']} !important;
        border-radius: 8px !important;
        color: {TH_COLORS['primary']} !important;
    }}
    
    .excel-bulk-section .stTextInput input:focus {{
        border-color: {TH_COLORS['primary']} !important;
        box-shadow: 0 0 0 2px {TH_COLORS['primary']}20 !important;
    }}
    
    .excel-bulk-section .stRadio {{
        background: white !important;
        border-radius: 8px !important;
        padding: 1rem !important;
        border: 1px solid {TH_COLORS['light_gray']} !important;
    }}
    
    .excel-bulk-section .stRadio * {{
        color: {TH_COLORS['primary']} !important;
        background: transparent !important;
    }}
    
    .excel-bulk-section .stFileUploader {{
        background: white !important;
        border: 2px dashed {TH_COLORS['primary']} !important;
        border-radius: 8px !important;
        padding: 2rem !important;
    }}
    
    .excel-bulk-section .stFileUploader * {{
        color: {TH_COLORS['primary']} !important;
        background: transparent !important;
    }}
    
    .excel-bulk-section .stFileUploader:hover {{
        border-color: {TH_COLORS['secondary']} !important;
        background: {TH_COLORS['light_gray']} !important;
    }}
    
    .excel-bulk-section .stButton button {{
        background: {TH_COLORS['primary']} !important;
        color: white !important;
        border-radius: 8px !important;
        padding: 0.75rem 2rem !important;
        font-weight: 600 !important;
        border: none !important;
    }}
    
    .excel-bulk-section .stButton button:hover {{
        background: {TH_COLORS['secondary']} !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.2) !important;
    }}
    
    .excel-bulk-section .stSuccess {{
        background: {TH_COLORS['primary']}10 !important;
        border: 1px solid {TH_COLORS['primary']} !important;
        border-radius: 8px !important;
        color: {TH_COLORS['primary']} !important;
    }}
    
    .excel-bulk-section .stInfo {{
        background: {TH_COLORS['accent']}10 !important;
        border: 1px solid {TH_COLORS['accent']} !important;
        border-radius: 8px !important;
        color: {TH_COLORS['secondary']} !important;
    }}
    
    .excel-bulk-section .stError {{
        background: #ffebee !important;
        border: 1px solid #f44336 !important;
        border-radius: 8px !important;
        color: #d32f2f !important;
    }}
    
    .excel-bulk-section .stWarning {{
        background: #fff3e0 !important;
        border: 1px solid #ff9800 !important;
        border-radius: 8px !important;
        color: #f57c00 !important;
    }}
    
    .excel-bulk-section .stExpander {{
        background: white !important;
        border: 1px solid {TH_COLORS['light_gray']} !important;
        border-radius: 8px !important;
    }}
    
    .excel-bulk-section .stExpander * {{
        color: {TH_COLORS['primary']} !important;
        background: transparent !important;
    }}
    
    .excel-bulk-section .stDataFrame {{
        background: white !important;
        border: 1px solid {TH_COLORS['light_gray']} !important;
        border-radius: 8px !important;
    }}
    
    .excel-bulk-section .stDataFrame * {{
        color: {TH_COLORS['primary']} !important;
        background: transparent !important;
    }}
    
    .excel-bulk-section .stMetric {{
        background: white !important;
        border: 1px solid {TH_COLORS['light_gray']} !important;
        border-radius: 8px !important;
        padding: 1rem !important;
    }}
    
    .excel-bulk-section .stMetric * {{
        color: {TH_COLORS['primary']} !important;
        background: transparent !important;
    }}
    
    /* Specific overrides for file uploader and dataframe */
    .excel-bulk-section [data-testid="stFileUploader"] {{
        background: white !important;
        border: 2px dashed {TH_COLORS['primary']} !important;
        border-radius: 8px !important;
        padding: 2rem !important;
    }}
    
    .excel-bulk-section [data-testid="stFileUploader"] * {{
        color: {TH_COLORS['primary']} !important;
        background: transparent !important;
    }}
    
    .excel-bulk-section [data-testid="stFileUploader"] .uploadedFile {{
        background: white !important;
        border: 1px solid {TH_COLORS['light_gray']} !important;
        color: {TH_COLORS['primary']} !important;
    }}
    
    .excel-bulk-section [data-testid="stDataFrame"] {{
        background: white !important;
        border: 1px solid {TH_COLORS['light_gray']} !important;
        border-radius: 8px !important;
    }}
    
    .excel-bulk-section [data-testid="stDataFrame"] * {{
        color: {TH_COLORS['primary']} !important;
        background: white !important;
    }}
    
    .excel-bulk-section [data-testid="stDataFrame"] table {{
        background: white !important;
        color: {TH_COLORS['primary']} !important;
    }}
    
    .excel-bulk-section [data-testid="stDataFrame"] th {{
        background: {TH_COLORS['light_gray']} !important;
        color: {TH_COLORS['primary']} !important;
        border: 1px solid {TH_COLORS['light_gray']} !important;
    }}
    
    .excel-bulk-section [data-testid="stDataFrame"] td {{
        background: white !important;
        color: {TH_COLORS['primary']} !important;
        border: 1px solid {TH_COLORS['light_gray']} !important;
    }}
    
    /* Override any Streamlit dark theme */
    .excel-bulk-section .stApp {{
        background: white !important;
    }}
    
    .excel-bulk-section .main {{
        background: white !important;
    }}
    
    .excel-bulk-section .block-container {{
        background: white !important;
    }}
</style>
""", unsafe_allow_html=True)

# Load assets
logo_bytes_default = get_embedded_logo()
sig_bytes_default = get_embedded_signature()

# Header
current_time = datetime.now().strftime("%B %d, %Y at %I:%M %p")

st.markdown(f"""
<div class="main-header">
    <div class="logo-container">
        {"<img src='data:image/png;base64,{}' alt='Total Health Logo'>".format(base64.b64encode(logo_bytes_default).decode()) if logo_bytes_default else "🏥"}
    </div>
    <h1>LOR/LOA Generator</h1>
    <div class="version-info">v46.0</div>
    <div class="status-badge">System Online</div>
    <div class="last-updated">Last updated: {current_time}</div>
</div>
""", unsafe_allow_html=True)

# Initialize session state
if 'document_type' not in st.session_state:
    st.session_state.document_type = 'LOR'
if 'novartis_clicked' not in st.session_state:
    st.session_state.novartis_clicked = False
if 'additional_info_text' not in st.session_state:
    st.session_state.additional_info_text = ""

# Main form
with st.container():
    st.markdown('<div class="stContainer">', unsafe_allow_html=True)
    
    # Document type selection
    st.markdown('<div class="section-header">Document Type</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📄 Letter of Request (LOR)", use_container_width=True):
            st.session_state.document_type = 'LOR'
            st.rerun()
    with col2:
        if st.button("📋 Letter of Agreement (LOA)", use_container_width=True):
            st.session_state.document_type = 'LOA'
            st.rerun()
    
    # Show selected document type
    if st.session_state.document_type == 'LOR':
        st.success("✅ Selected: Letter of Request (LOR)")
    else:
        st.success("✅ Selected: Letter of Agreement (LOA)")
    
    st.markdown('</div>', unsafe_allow_html=True)

# Event and company information
with st.container():
    st.markdown('<div class="stContainer">', unsafe_allow_html=True)
    st.markdown('<div class="section-header">Event & Company Information</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # Event search and selection
        st.markdown("**Search Events:**")
        event_search = st.text_input("🔍 Search events by name, year, or location", key="event_search", placeholder="e.g., '2025', 'Memphis', 'ASCO'")
        
        # Get upcoming events (chronologically sorted, past events auto-dropped)
        upcoming_events = get_upcoming_events()
        
        # Filter events based on search
        if event_search:
            filtered_events = [e for e in upcoming_events if event_search.lower() in e["meeting_name"].lower() or 
                              event_search.lower() in e["city_state"].lower() or 
                              event_search.lower() in e["venue"].lower()]
        else:
            # Use chronologically sorted upcoming events
            filtered_events = upcoming_events
        
        if not filtered_events:
            st.warning("🔍 No upcoming events found matching your search.")
            st.info("💡 **Tip:** Past events are automatically filtered out. Try searching for a different term or check back for new events.")
            filtered_events = []
        else:
            # Show count of upcoming events
            total_upcoming = len(upcoming_events)
            filtered_count = len(filtered_events)
            if event_search:
                st.success(f"📅 **Found {filtered_count} upcoming events** matching your search (out of {total_upcoming} total upcoming events)")
            else:
                st.success(f"📅 **Showing {total_upcoming} upcoming events** in chronological order")
        
        # Event selection dropdown
        if filtered_events:
            event_names = [e["meeting_name"] for e in filtered_events]
            event_choice = st.selectbox("Select Event", event_names, key="event_select")
            event = next(e for e in filtered_events if e["meeting_name"] == event_choice)
        else:
            st.error("❌ **No events available** - Please try a different search or contact support.")
            event_choice = None
            event = None
        
        # Show selected event with more details
        if event_choice and event:
            st.info(f"📅 **Selected Event:** {event_choice}")
            st.caption(f"📍 **Location:** {event['venue']}, {event['city_state']}")
            st.caption(f"📅 **Date:** {event['meeting_date_long']}")
            
            # Check if this is a past event and warn user
            try:
                event_date = parse_event_date(event['meeting_date_long'])
                today = datetime.now()
                if event_date < today:
                    st.warning("⚠️ **PAST EVENT WARNING:** This event has already occurred. You can still generate letters for past events if needed.")
            except:
                pass  # If date parsing fails, don't show warning
            
            # Show event type and compliance info
            event_year = get_event_year(event_choice)
            if "ASCO" in event_choice or "ESMO" in event_choice:
                st.caption("🏥 **Event Type:** CME-Accredited Educational Event")
            elif "Best of" in event_choice:
                st.caption("🏥 **Event Type:** Premium Educational Conference")
            else:
                st.caption("🏥 **Event Type:** Educational Conference")
        else:
            st.warning("⚠️ **Please select an event** to continue")
        
        # Company information
        st.markdown("**Company Information:**")
        company_name = st.text_input(
            "Company Name (Exhibitor)", 
            key="company_name",
            placeholder="e.g., Novartis Pharmaceuticals Corporation",
            help="Enter the full legal name of the sponsoring company"
        )
        
        # LOA-specific fields
        if st.session_state.document_type == 'LOA':
            company_address = st.text_area(
                "Company Address (for LOA)", 
                key="company_address",
                placeholder="Enter full company address for the agreement",
                help="Full legal address of the sponsoring company"
            )
            
            agreement_date = st.date_input(
                "Agreement Date (for LOA)", 
                value=dt.date.today(),
                key="agreement_date",
                help="Date when the agreement is being made"
            )
            
            signature_person = st.selectbox(
                "Select Signatory for LOA:",
                options=[
                    "Sarah Louden - Founder and Executive Director, Total Health Conferencing",
                    "Michael Eisinger - Manager, Business Operations"
                ],
                key="signature_person",
                help="Choose who will sign the LOA"
            )
        
        attendance = st.number_input(
            "Expected Attendance (optional)", 
            min_value=0, 
            step=10, 
            value=0, 
            key="attendance",
            help="Expected number of attendees. Leave as 0 if not specified. Typical ranges: 50-200 for regional events, 200-500 for national events."
        )
        
        if attendance > 0:
            st.caption(f"📊 **Attendance Note:** {attendance} expected attendees")
        else:
            st.caption("📊 **Attendance Note:** Will use general attendance language")
    
    with col2:
        st.markdown("**Signer:** Sarah Louden (fixed)")
        st.caption("Using 'th_logo.(png/jpg)' and 'sarah_signature.(png/jpg)' from this folder.")
    
    st.markdown('</div>', unsafe_allow_html=True)

# Booth and add-ons
with st.container():
    st.markdown('<div class="stContainer">', unsafe_allow_html=True)
    st.markdown('<div class="section-header">Booth & Add-Ons</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Booth selection with enhanced guidance
        st.markdown("**Exhibit Booth Selection:**")
        booth_options = ["(no booth)"] + list(BOOTH_PRICES.keys())
        booth_choice = st.selectbox(
            "Booth Tier (optional)",
            options=booth_options,
            index=0,
            format_func=lambda k: {
                "(no booth)": "(No Booth)",
                "standard_1d": f"Standard (1 day) - ${BOOTH_PRICES['standard_1d']:,}",
                "standard_2d": f"Standard (2 days) - ${BOOTH_PRICES['standard_2d']:,}",
                "platinum": f"Platinum - ${BOOTH_PRICES['platinum']:,}",
                "best_of": f"Best of - ${BOOTH_PRICES['best_of']:,}",
                "premier": f"Premier - ${BOOTH_PRICES['premier']:,}",
            }[k],
            key="booth_select",
            help="Select booth tier based on your sponsorship level. Standard booths include basic exhibit space, Platinum includes enhanced placement, Best of/Premier include premium positioning."
        )
        booth_selected = booth_choice != "(no booth)"
        
        if booth_selected:
            # Get the formatted display name for the selected booth
            booth_display_names = {
                "standard_1d": "Standard (1 day)",
                "standard_2d": "Standard (2 days)", 
                "platinum": "Platinum",
                "best_of": "Best of",
                "premier": "Premier"
            }
            booth_display_name = booth_display_names.get(booth_choice, booth_choice)
            st.success(f"✅ **Selected Booth:** {booth_display_name}")
            st.caption("📋 **Booth includes:** 6' draped table, 2 chairs, 2 badges, logo placement, bag insert, attendee list")
        else:
            st.info("ℹ️ **No booth**")
            st.caption("📋 **Educational support:** Logo placement, virtual presence, attendee list")
    
    with col2:
        # Get pricing based on event year
        event_year = get_event_year(event_choice)
        add_ons_pricing = get_add_ons_pricing(event_year)
        
        # Organize add-ons by category
        add_on_categories = {
            "📱 Digital Marketing": ["platform_banner", "email_banner", "registration_banner"],
            "📄 Print Materials": ["program_ad_full"],
            "🔌 On-Site Services": ["charging_stations", "wifi_sponsorship"],
            "🤝 Networking": ["networking_reception", "networking_activity"],
            "🎓 Educational Programs": ["advisory_board", "non_cme_session"]
        }
        
        st.markdown("**Add-Ons by Category:**")
        
        # Create expandable sections for each category
        selected_add_ons = []
        
        for category, add_on_list in add_on_categories.items():
            with st.expander(f"{category} ({len([a for a in add_on_list if a in add_ons_pricing])} options)"):
                category_add_ons = []
                for add_on_key in add_on_list:
                    if add_on_key in add_ons_pricing:
                        if st.checkbox(
                            f"{add_ons_pricing[add_on_key]['label']} - {currency(add_ons_pricing[add_on_key]['price'])}",
                            key=f"addon_{add_on_key}",
                            help=ADD_ON_BULLETS.get(add_on_key, ["No description available"])[0] if ADD_ON_BULLETS.get(add_on_key) else "No description available"
                        ):
                            category_add_ons.append(add_on_key)
                selected_add_ons.extend(category_add_ons)
        
        # Update the add_on_keys variable
        add_on_keys = selected_add_ons
        
        if add_on_keys:
            st.success(f"✅ **Selected {len(add_on_keys)} add-on(s)**")
            total_addon_cost = sum(add_ons_pricing[k]["price"] for k in add_on_keys)
            st.caption(f"💰 **Add-ons total:** {currency(total_addon_cost)}")
        else:
            st.info("ℹ️ **No add-ons selected**")
    
    st.markdown('</div>', unsafe_allow_html=True)

# Additional information with Novartis button
with st.container():
    st.markdown('<div class="stContainer">', unsafe_allow_html=True)
    st.markdown('<div class="section-header">Additional Information & Compliance</div>', unsafe_allow_html=True)
    
    # Compliance guidance
    st.markdown("""
    **📋 Compliance Guidelines:**
    - **ACCME Standards:** All educational content must maintain independence from commercial influence
    - **PhRMA Code:** Ensure proper separation between educational and promotional activities  
    - **OIG Guidance:** Maintain fair market value for all transactions
    - **Sunshine Act:** All payments will be reported as required by federal and state laws
    """)
    
    # Document type specific guidance
    if st.session_state.document_type == 'LOR':
        st.info("📄 **Letter of Request (LOR):** Use for educational/CME support requests. Focus on educational value and compliance with ACCME standards.")
    else:
        st.info("📋 **Letter of Agreement (LOA):** Use for commercial agreements. Include specific terms, deliverables, and compliance requirements.")
    
    # Novartis standard text
    NOVARTIS_TEXT = """Novartis Pharmaceuticals Corporation (NPC) agrees to participate in this event according to the following terms and conditions:
• The arrangement to participate in this event represents a fair market value transaction for bona fide services rendered to NPC and is specifically made without intent to induce or reward referrals for the purchase of NPC products.
• The payment is intended to cover only the costs related to the agreed upon promotional activities.
• NPC funds will not be used for items that are not permitted by applicable industry code requirements, and will not be used to remunerate any individual healthcare professional (HCP), nor will any funds be used to support any Independent Charitable Co-Pay Assistance programs.
• The exhibit payment is not an unrestricted charitable contribution or educational grant for accredited or non-accredited continuing medical education.
• This payment will be reported where applicable to meet Federal and State Sunshine Disclosure laws."""
    
    # Compliance templates section
    st.markdown("**📋 Compliance Templates:**")
    
    # Template buttons with proper session state handling
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("🏢 Novartis Standard", use_container_width=True):
            st.session_state.template_text = NOVARTIS_TEXT
            st.session_state.template_name = "Novartis Standard"
            st.rerun()
    
    with col2:
        if st.button("🏥 Generic Pharma", use_container_width=True):
            # Get the company name from the form
            company_name = st.session_state.get('company_name', '[Company Name]')
            st.session_state.template_text = f"""{company_name} agrees to participate in this event according to the following terms and conditions:
• This arrangement represents a fair market value transaction for bona fide services rendered.
• The payment is intended to cover only the costs related to the agreed upon activities.
• Funds will not be used for items not permitted by applicable industry code requirements.
• This payment will be reported where applicable to meet Federal and State Sunshine Disclosure laws."""
            st.session_state.template_name = "Generic Pharma"
            st.rerun()
    
    with col3:
        if st.button("🎓 Educational Focus", use_container_width=True):
            st.session_state.template_text = """This educational support is provided to advance medical knowledge and improve patient care:
• Educational content will maintain independence from commercial influence.
• All activities will comply with ACCME Standards for Commercial Support.
• The educational value and learning objectives will be clearly defined.
• Content will be evidence-based and free from commercial bias."""
            st.session_state.template_name = "Educational Focus"
            st.rerun()
    
    # BULLETPROOF TEMPLATE SYSTEM - No key conflicts!
    # Initialize session state for additional info text
    if 'additional_info_text' not in st.session_state:
        st.session_state.additional_info_text = ''
    
    # Get template text from session state
    template_text = st.session_state.get('template_text', '')
    
    # If template was clicked, update the additional info text
    if template_text and 'template_name' in st.session_state:
        st.session_state.additional_info_text = template_text
        # Clear the template text to prevent re-setting
        st.session_state.template_text = ''
    
    # Text area for additional information - NO KEY TO AVOID CONFLICTS!
    company_required_text = st.text_area(
        "Additional Information (optional)",
        value=st.session_state.additional_info_text,
        placeholder="Paste company-required compliance text here. First non-empty line = lead-in; each subsequent line becomes a bullet.",
        height=150,
        help="Enter company-specific compliance language. The first line will be used as a lead-in paragraph, and subsequent lines will be formatted as bullet points. Use the compliance templates above for common scenarios."
    )
    
    # Update session state with current text
    st.session_state.additional_info_text = company_required_text
    
    # Show template loaded message
    if 'template_name' in st.session_state:
        st.success(f"✅ {st.session_state.template_name} template loaded")
        # Clear the template name to prevent re-showing
        del st.session_state.template_name
    
    st.markdown('</div>', unsafe_allow_html=True)

# Discount and pricing
with st.container():
    st.markdown('<div class="stContainer">', unsafe_allow_html=True)
    st.markdown('<div class="section-header">Pricing & Discounts</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**Pricing & Discounts:**")
        discount_key = st.selectbox(
            "Discount",
            options=list(DISCOUNT_MAP.keys()),
            index=0,
            format_func=lambda k: {
                "none": "None",
                "minus_10": "-10%",
                "minus_15": "-15%",
                "minus_20": "-20%",
                "custom": "Custom Total Override",
            }[k],
            key="discount_select",
            help="Select appropriate discount based on partnership level and timing. Custom override allows for special pricing arrangements."
        )
        
        if discount_key == "custom":
            custom_total = st.number_input(
                "Custom Total ($)", 
                min_value=0, 
                step=50, 
                key="custom_total",
                help="Enter the final agreed-upon total amount. This will override all calculated pricing."
            )
        else:
            custom_total = None
    
    with col2:
        # Calculate pricing
        base = BOOTH_PRICES.get(booth_choice, 0)
        addons_total = sum(add_ons_pricing[k]["price"] for k in add_on_keys)
        subtotal = base + addons_total
        
        if discount_key == "custom" and custom_total:
            final_total = custom_total
        else:
            factor = DISCOUNT_MAP[discount_key] or 1.0
            final_total = round_nearest_50(subtotal * factor)
        
        st.markdown("### 💰 Pricing Summary")
        
        # Create a more detailed pricing breakdown
        pricing_data = {
            "Base Booth": currency(base) if booth_selected else "No booth selected",
            "Add-ons": currency(addons_total) if add_on_keys else "No add-ons selected",
            "Subtotal": currency(subtotal),
        }
        
        if discount_key == "custom" and custom_total:
            pricing_data["Custom Override"] = currency(custom_total)
            pricing_data["Final Total"] = currency(final_total)
        else:
            discount_percent = {"none": "0%", "minus_10": "10%", "minus_15": "15%", "minus_20": "20%", "custom": "custom"}[discount_key]
            if discount_percent != "0%":
                discount_amount = subtotal - final_total
                pricing_data[f"Discount ({discount_percent})"] = f"-{currency(discount_amount)}"
            pricing_data["Final Total"] = currency(final_total)
        
        # Display pricing breakdown
        for item, amount in pricing_data.items():
            if item == "Final Total":
                st.markdown(f"**{item}:** {amount}")
            else:
                st.write(f"**{item}:** {amount}")
        
        # Add compliance note
        st.caption("💡 **Note:** All pricing represents fair market value and will be reported as required by Sunshine Act regulations.")
    
    st.markdown('</div>', unsafe_allow_html=True)

# Progress indicator
st.markdown("---")
st.markdown("### 📋 Document Generation Progress")

# Check if all required fields are filled
required_fields = {
    "Event": event_choice if event_choice else None,
    "Company": company_name,
    "Document Type": st.session_state.document_type
}

missing_fields = [field for field, value in required_fields.items() if not value or value.strip() == ""]

if missing_fields:
    st.warning(f"⚠️ **Please complete:** {', '.join(missing_fields)}")
    st.info("💡 **Tip:** All required fields must be completed before generating the document.")
else:
    st.success("✅ **All required fields completed** - Ready to generate document!")

st.markdown("---")

# Generate button and downloads
button_text = "🚀 Generate LOA" if st.session_state.document_type == 'LOA' else "🚀 Generate Letter"
if st.button(button_text, use_container_width=True):
    payload = {
        "company_name": (company_name or "[Company]").strip(),
        "meeting_name": event["meeting_name"],
        "meeting_date_long": event["meeting_date_long"],
        "venue": event["venue"],
        "city_state": event["city_state"],
        "attendance_expected": attendance or None,
        "amount_currency": currency(final_total),
    }
    
    # Add LOA-specific fields if LOA is selected
    if st.session_state.document_type == 'LOA':
        # Get booth tier and price from booth_choice
        booth_tier = booth_choice if booth_selected else None
        booth_price = BOOTH_PRICES.get(booth_choice, 0) if booth_selected else 0
        
        payload.update({
            "company_address": st.session_state.get("company_address", "[Address]"),
            "agreement_date": st.session_state.get("agreement_date", dt.date.today().strftime("%B %d, %Y")),
            "signature_person": st.session_state.get("signature_person", "Sarah Louden - Founder and Executive Director, Total Health Conferencing"),
            "booth_tier": booth_tier,
            "booth_price": currency(booth_price) if booth_selected else None,
            "additional_info": company_required_text
        })
    
    if st.session_state.document_type == 'LOA':
        paragraphs = render_loa_paragraphs(payload, booth_selected, add_on_keys, add_ons_pricing)
    else:
        paragraphs = render_letter_paragraphs(payload, st.session_state.document_type)
    
    # Preview
    preview_lines = "\n".join(paragraphs) + "\n\n"
    if booth_selected:
        preview_lines += "Exhibit Booth\n"
        preview_lines += "\n".join([f"• {b}" for b in booth_bullets()]) + "\n\n"
    if add_on_keys:
        preview_lines += "Selected Add-Ons\n"
        for k in add_on_keys:
            preview_lines += f"{add_ons_pricing[k]['label']}\n"
            preview_lines += "\n".join([f"• {b}" for b in ADD_ON_BULLETS.get(k, [])]) + "\n\n"
    if company_required_text:
        lead_in, bullets = parse_additional_info(company_required_text)
        preview_lines += "Additional Information\n"
        if lead_in:
            preview_lines += lead_in + "\n"
        if bullets:
            preview_lines += "\n".join([f"• {b}" for b in bullets]) + "\n"
        preview_lines += "\n"
    preview_lines += "Grateful for your support,\n\n"
    preview_lines += f"{SARAH['name']}\n{SARAH['title']}\n{SARAH['email']}\n{SARAH['phone']}"
    
    st.text_area("📄 Document Preview", preview_lines, height=400)
    
    base_filename = f"{st.session_state.document_type}_{(company_name or 'Company').replace(' ', '_')}_{event['meeting_name'].replace(' ', '_')}"
    
    logo_bytes = logo_bytes_default
    sig_bytes = sig_bytes_default
    
    # DOCX Download
    if st.session_state.document_type == 'LOA':
        docx_bytes = build_loa_docx_bytes(paragraphs, logo_bytes, sig_bytes)
        st.download_button(
            "📄 Download LOA (.docx)", 
            data=docx_bytes,
            file_name=base_filename + ".docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if DOCX_AVAILABLE else "text/plain"
        )
    else:
        docx_bytes = build_docx_bytes(paragraphs, booth_selected, add_on_keys, logo_bytes, sig_bytes, company_required_text)
        st.download_button(
            "📄 Download Letter (.docx)", 
            data=docx_bytes,
            file_name=base_filename + ".docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if DOCX_AVAILABLE else "text/plain"
        )
    
    # PDF Download
    if st.session_state.document_type == 'LOA':
        pdf_bytes = build_loa_pdf_bytes(paragraphs, logo_bytes, sig_bytes)
        st.download_button(
            "📄 Download LOA (.pdf)", 
            data=pdf_bytes,
            file_name=base_filename + ".pdf",
            mime="application/pdf" if PDF_AVAILABLE else "text/plain"
        )
    else:
        pdf_bytes = build_pdf_bytes(paragraphs, booth_selected, add_on_keys, logo_bytes, sig_bytes, company_required_text)
        st.download_button(
            "📄 Download Letter (.pdf)", 
            data=pdf_bytes,
            file_name=base_filename + ".pdf",
            mime="application/pdf" if PDF_AVAILABLE else "text/plain"
        )


# ======================== EXCEL BULK GENERATION ========================
st.markdown("---")
st.markdown("---")

# Excel Bulk Generation Tab
st.markdown("### 📊 Excel Bulk Letter Generator")
st.markdown("**Upload your exhibitor invite spreadsheet → Download ZIP with all letters**")


# Excel bulk section - clean approach without CSS overrides

# Instructions
with st.expander("📋 How To Use Excel Bulk Mode", expanded=False):
    st.markdown("""
    ### Step-by-Step Guide:
    
    1. **Prepare your Excel file** with these columns:
       - `Exhibitor Invite` (required) - Used for filename
       - `Event Name` (required) - Will be matched to system events  
       - `Total` (required) - Final price shown in letter
       - `Company Name` (optional) - Company name in letter (defaults to "[Company Name]")
       - `Expected Attendance` (optional) - Override system attendance
       - `Date`, `City`, `Venue`, `Official Address` (optional) - Override system defaults
       - `Amount`, `Discount` (optional) - For reference
    
    2. **Upload the Excel file** below
    
    3. **Review the preview** to make sure data loaded correctly
    
    4. **Click Generate** button
    
    5. **Download ZIP** with all your letters!
    
    ### Example Excel Format:
    ```
    Exhibitor Invite | Event Name | Date | City | Venue | Official Address | Amount | Discount | Total
    2026 ASCO Denver Invite | 2026 ASCO Direct Denver | June 27-28 | Denver, CO | Denver Marriott Westminster | 7000 Church Ranch Blvd... | $7,500 | 10% | $6,750
    ```
    """)

# Company name input (applies to all letters)
st.markdown("#### 🏢 Company Information")
bulk_company_name = st.text_input(
    "Company Name (applies to all letters in Excel)",
    placeholder="e.g., Novartis Pharmaceuticals Corporation",
    help="This company name will be used in all generated letters",
    key="bulk_company_name"
)

if bulk_company_name:
    st.success(f"✅ Company: {bulk_company_name}")
else:
    st.info("💡 Enter company name above - it will be used for all letters")

# Document type selection for bulk
bulk_doc_type = st.radio(
    "Select Document Type for Bulk Generation",
    options=['LOR', 'LOA'],
    index=0,
    horizontal=True,
    help="LOR = Letter of Request, LOA = Letter of Agreement",
    key="bulk_doc_type"
)

st.markdown("---")

# File uploader for Excel - clean approach
st.markdown("#### 📁 Upload Excel File")
st.markdown("Upload your exhibitor invite spreadsheet (.xlsx or .xls)")

uploaded_file = st.file_uploader(
    "Choose Excel file",
    type=['xlsx', 'xls'],
    help="Upload your exhibitor invite spreadsheet",
    key="bulk_excel_uploader"
)

if uploaded_file:
    try:
        # Import pandas for Excel processing
        import pandas as pd
        import zipfile
        import re
        
        # Read Excel
        df = pd.read_excel(uploaded_file)
        
        st.success(f"✅ **File loaded successfully:** {len(df)} rows found")
        
        # Show preview
        st.markdown("#### 📊 Data Preview (first 10 rows)")
        st.dataframe(df.head(10), use_container_width=True)
        
        # Validate columns
        required_cols = ['Exhibitor Invite', 'Event Name', 'Total']
        missing_cols = [col for col in required_cols if col not in df.columns]
        
        if missing_cols:
            st.error(f"❌ **Missing required columns:** {', '.join(missing_cols)}")
            st.info("💡 Your Excel must have at least: 'Exhibitor Invite', 'Event Name', and 'Total' columns")
        else:
            st.success(f"✅ **All required columns found!**")
            st.info(f"📋 **Columns in your file:** {', '.join(df.columns.tolist())}")
            
            # Generate button
            if st.button(f"🚀 Generate {len(df)} {bulk_doc_type}s from Excel", use_container_width=True, key="bulk_generate_btn"):
                
                with st.spinner(f"Generating {len(df)} documents..."):
                    zip_buffer = BytesIO()
                    generated_count = 0
                    error_count = 0
                    errors = []
                    
                    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                        for idx, row in df.iterrows():
                            try:
                                # Extract data
                                exhibitor_invite = str(row.get('Exhibitor Invite', f'Invite_{idx}'))
                                event_name = str(row.get('Event Name', ''))
                                
                                # Use company name from text input first, then Excel column, then default
                                if bulk_company_name and bulk_company_name.strip():
                                    company_name = bulk_company_name.strip()
                                elif 'Company Name' in row and str(row.get('Company Name')) != 'nan':
                                    company_name = str(row.get('Company Name'))
                                else:
                                    company_name = '[Company Name]'
                                
                                date_str = str(row.get('Date', ''))
                                city = str(row.get('City', ''))
                                venue = str(row.get('Venue', ''))
                                total_str = str(row.get('Total', '$0'))
                                
                                # Parse amount
                                total = float(re.sub(r'[$,]', '', total_str))
                                
                                # Parse expected attendance if provided
                                expected_attendance = None
                                if 'Expected Attendance' in row:
                                    try:
                                        att_val = row.get('Expected Attendance')
                                        if att_val and str(att_val).strip() and str(att_val) != 'nan':
                                            expected_attendance = int(float(att_val))
                                    except:
                                        pass
                                
                                # Find matching event with improved fuzzy matching
                                matching_event = None
                                
                                # Normalize event name for better matching
                                def normalize_name(name):
                                    """Normalize event name for fuzzy matching"""
                                    name = name.lower()
                                    # Replace various dashes and special chars
                                    name = name.replace('–', '-').replace('—', '-').replace('−', '-')
                                    # Remove common variations
                                    name = name.replace('best of asco', 'asco direct')
                                    name = name.replace('best of', '').replace('best', '')
                                    # Remove extra spaces
                                    name = ' '.join(name.split())
                                    return name
                                
                                normalized_excel = normalize_name(event_name)
                                
                                # Try exact partial match first
                                for evt in EVENTS:
                                    evt_name_lower = evt['meeting_name'].lower()
                                    if event_name.lower() in evt_name_lower or evt_name_lower in event_name.lower():
                                        matching_event = evt
                                        break
                                
                                # If not found, try normalized matching
                                if not matching_event:
                                    for evt in EVENTS:
                                        normalized_system = normalize_name(evt['meeting_name'])
                                        if normalized_excel in normalized_system or normalized_system in normalized_excel:
                                            matching_event = evt
                                            break
                                
                                # If still not found, try key word matching (year + location)
                                if not matching_event:
                                    excel_words = set(event_name.lower().split())
                                    for evt in EVENTS:
                                        evt_words = set(evt['meeting_name'].lower().split())
                                        # Match if they share year and location/type keywords
                                        common = excel_words & evt_words
                                        if len(common) >= 3:  # At least 3 words match
                                            matching_event = evt
                                            break
                                
                                if not matching_event:
                                    errors.append(f"Row {idx+2}: Event '{event_name}' not found in system")
                                    error_count += 1
                                    continue
                                
                                # Get attendance - use Excel column if provided, otherwise use system event data
                                if expected_attendance:
                                    final_attendance = expected_attendance
                                else:
                                    final_attendance = matching_event.get('expected_attendance', None)
                                
                                # Create payload
                                payload = {
                                    "company_name": company_name if company_name and company_name != 'nan' else "[Company Name]",
                                    "meeting_name": matching_event["meeting_name"],
                                    "meeting_date_long": date_str if date_str and date_str != 'nan' else matching_event["meeting_date_long"],
                                    "venue": venue if venue and venue != 'nan' else matching_event["venue"],
                                    "city_state": city if city and city != 'nan' else matching_event["city_state"],
                                    "attendance_expected": final_attendance,
                                    "amount_currency": currency(total),
                                }
                                
                                # Generate paragraphs
                                paragraphs = render_letter_paragraphs(payload, bulk_doc_type)
                                
                                # Load assets
                                logo_bytes = get_embedded_logo()
                                sig_bytes = get_embedded_signature()
                                
                                # Generate documents
                                docx_bytes = build_docx_bytes(paragraphs, True, [], logo_bytes, sig_bytes, "")
                                pdf_bytes = build_pdf_bytes(paragraphs, True, [], logo_bytes, sig_bytes, "")
                                
                                # Clean filename
                                safe_filename = "".join(c for c in exhibitor_invite if c.isalnum() or c in (' ', '-', '_')).rstrip()
                                safe_filename = safe_filename.replace('  ', ' ').replace(' ', '_')[:80]
                                
                                # Add to ZIP in organized folders
                                zip_file.writestr(f"Word_Documents/{safe_filename}.docx", docx_bytes)
                                zip_file.writestr(f"PDF_Documents/{safe_filename}.pdf", pdf_bytes)
                                generated_count += 1
                                
                            except Exception as e:
                                errors.append(f"Row {idx+2}: {str(e)}")
                                error_count += 1
                    
                    # Prepare ZIP for download
                    zip_buffer.seek(0)
                    
                    # Create clean company name for filename
                    if bulk_company_name and bulk_company_name.strip():
                        safe_company = "".join(c for c in bulk_company_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
                        safe_company = safe_company.replace('  ', ' ').replace(' ', '_')[:40]
                        zip_filename = f"{bulk_doc_type}_{safe_company}_{len(df)}_Letters_{dt.date.today().strftime('%Y%m%d')}.zip"
                    else:
                        zip_filename = f"{bulk_doc_type}_Excel_Bulk_{len(df)}_Letters_{dt.date.today().strftime('%Y%m%d')}.zip"
                    
                    # Download button
                    st.download_button(
                        f"📦 Download ZIP with {generated_count} {bulk_doc_type}s",
                        data=zip_buffer.getvalue(),
                        file_name=zip_filename,
                        mime="application/zip",
                        use_container_width=True,
                        key="bulk_download_btn"
                    )
                    
                    # Results
                    if generated_count > 0:
                        st.success(f"🎉 **Successfully Generated:** {generated_count} documents ({generated_count * 2} files total)")
                    
                    if error_count > 0:
                        st.warning(f"⚠️ **Errors:** {error_count} rows skipped due to issues")
                        with st.expander("📋 View Error Details"):
                            for error in errors:
                                st.write(f"• {error}")
                    
                    # Summary
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("✅ Succeeded", generated_count)
                    with col2:
                        st.metric("❌ Failed", error_count)
                    with col3:
                        st.metric("📊 Total Rows", len(df))
    
    except Exception as e:
        st.error(f"❌ **Error reading Excel file:** {str(e)}")
        st.info("💡 Make sure your file is a valid Excel format (.xlsx or .xls)")

else:
    st.info("👆 **Upload an Excel file to get started**")
    
    # Example data - clean approach
    st.markdown("#### 📋 Example Excel Data")
    st.markdown("Create your Excel file with the same column structure as shown below:")
    
    # Create example dataframe
    example_data = {
        'Exhibitor Invite': ['2026 ASCO Denver Invite', '2026 ESMO West Invite'],
        'Event Name': ['2026 ASCO Direct Denver', '2026 ESMO USA West'],
        'Date': ['June 27-28', 'November 14-15'],
        'City': ['Denver, CO', 'Colorado Springs, CO'],
        'Total': ['$6,750', '$6,750']
    }
    
    try:
        import pandas as pd
        example_df = pd.DataFrame(example_data)
        st.dataframe(example_df, use_container_width=True)
        
    except:
        st.write("Example columns: Exhibitor Invite, Event Name, Date, City, Total")
    
    st.info("💡 **Tip:** Create your Excel file with the same column structure as shown above")

st.markdown("---")

# ======================== ADDITIONAL RESOURCES ========================
st.markdown("---")
st.markdown("---")

with st.expander("📚 Additional Resources & Help", expanded=False):
    st.markdown("### 📚 Help & Documentation")
    st.info("""
    **All Features Integrated in This Application:**
    
    ✅ **Single Event Mode** - Generate individual LOR/LOA documents
    ✅ **Multi-Meeting Mode** - Generate quarterly package documents  
    ✅ **Excel Bulk Mode** - Upload Excel spreadsheet for bulk document generation
    
    **Need Help?**
    - All functionality is available in this single application
    - Use the Excel Bulk Letter Generator section above for bulk processing
    - Contact your system administrator for technical support
    """)

